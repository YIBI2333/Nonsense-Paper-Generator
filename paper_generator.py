#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
胡说八道论文生成器
生成形式严谨、内容胡说八道的学术论文，支持多种论文类型和参数调节

功能特性：
- 支持关键词输入（0-3个，不填则随机）
- 支持4种论文类型：学术期刊论文、硕士学位论文、博士学位论文、会议论文
- 支持胡说八道程度调节（1-10级）
- 支持3种篇幅选择：短篇、中篇、长篇
- 包含完整的学术论文结构：摘要、关键词、正文、表格、图表、公式、参考文献、致谢
- 导出Word格式（.docx）
"""

import random
import string
import argparse
import os
import tempfile
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont


# ============== 配置常量 ==============

# 字体配置
FONT_BODY_CN = 'SimSun'
FONT_BODY_EN = 'Times New Roman'
FONT_HEADING_CN = 'SimHei'
FONT_HEADING_EN = 'Arial'
FONT_FORMULA = 'Cambria Math'

# 字体中文名称（用于显示）
FONT_BODY_CN_NAME = '宋体'
FONT_HEADING_CN_NAME = '黑体'

# 页面配置
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_CM = 2.5

# 正文字号
BODY_FONT_SIZE = 12
HEADING1_SIZE = 16
HEADING2_SIZE = 14
HEADING3_SIZE = 12
CAPTION_SIZE = 10.5
ABSTRACT_KEYWORD_SIZE = 12


# ============== 数据池 ==============

# 关键词池
KEYWORD_POOLS = {
    'physics': [
        '量子纠缠', '暗物质', '引力波', '弦理论', '黑洞', '虫洞',
        '量子隧穿', '薛定谔的猫', '海森堡不确定性', '玻色-爱因斯坦凝聚',
        '拓扑相变', '量子涨落', '真空极化', '希格斯场', '中微子振荡',
        '夸克禁闭', '渐近自由', '规范对称性', '自发对称性破缺', '重整化群'
    ],
    'food': [
        '奶茶', '螺蛳粉', '煎饼果子', '烤红薯', '火锅', '烧烤',
        '麻辣烫', '小龙虾', '臭豆腐', '糖葫芦', '棉花糖', '爆米花',
        '薯条', '汉堡', '披萨', '寿司', '拉面', '饺子',
        '包子', '油条'
    ],
    'tech': [
        '区块链', '元宇宙', '人工智能', '大数据', '云计算', '物联网',
        '5G', '量子计算', '神经网络', '深度学习', '增强现实', '虚拟现实',
        '边缘计算', '数字孪生', '联邦学习', '知识图谱', '自然语言处理',
        '计算机视觉', '机器人', '自动驾驶'
    ],
    'social': [
        '内卷', '躺平', '摸鱼', '社恐', '社牛', 'emo',
        '破防', '绝绝子', 'yyds', 'awsl', '打工人', '干饭人',
        '尾款人', '韭菜', '吃瓜', '凡尔赛', 'PUA', '996',
        '007', '佛系'
    ],
    'academic': [
        '热力学', '混沌理论', '统计力学', '流体力学', '固体力学',
        '量子力学', '经典力学', '电动力学', '相对论', '量子场论',
        '凝聚态物理', '粒子物理', '天体物理', '宇宙学', '弦理论',
        'M理论', '圈量子引力', '因果动力学三角剖分', '熵增原理',
        '最小作用量原理'
    ]
}

# 作者名字
AUTHOR_NAMES = [
    '王德发', '张无序', '李混沌', '王量子', '刘熵增', '陈涨落',
    '杨隧穿', '赵相干', '孙叠加', '周纠缠', '吴跃迁', '郑散射',
    '王衍射', '冯干涉', '陈偏振', '褚共振', '卫色散', '蒋耗散',
    '沈涨落', '韩临界', '杨相变', '朱分岔', '秦分形', '尤混沌',
    '许自组织', '何涌现', '吕对称性', '施破缺', '张重整化', '孔正规化',
    '曹微扰', '严非微扰', '华拓扑', '金曲率', '魏挠率', '陶联络',
    '姜纤维丛', '戚流形', '谢同伦', '邹同调', '喻上同调', '柏范畴',
    '水函子', '窦自然变换', '章米田引理', '云伴随函子', '苏单子',
    '潘代数', '葛余代数', '奚双代数', '范Hopf代数', '彭李代数',
    '郎李群', '鲁表示论', '韦特征标', '马权代数', '苗顶点算子代数'
]

# 大学/机构
UNIVERSITIES = [
    '非线性早餐研究所', '宇宙袜子失踪现象研究中心', '平行时空脆度实验室',
    '量子奶茶研究中心', '螺蛳粉混沌理论实验室', '煎饼果子统计力学研究所',
    '烤红薯引力波观测站', '算命神经网络研究院', '摸鱼动力学研究中心',
    '内卷热力学实验室', '躺平量子态研究所', '社恐场论研究中心',
    '干饭人热力学协会', '韭菜生长动力学实验室', '凡尔赛统计研究中心',
    '996熵增研究所', '007相变研究中心', '佛系量子态实验室',
    '破防临界现象研究中心', 'emo涨落动力学实验室', '绝绝子共振研究所',
    'yyds对称性破缺研究中心', 'awsl量子隧穿实验室', '吃瓜拓扑场论研究所'
]

# 期刊名
JOURNALS = [
    '《胡说物理学报》', '《混沌与无序》', '《量子早餐动力学》',
    '《应用胡说八道》', '《理论扯淡》', '《实验废话》',
    '《现代伪科学》', '《虚构物理评论》', '《抽象研究快报》',
    '《非线性扯淡》', '《统计胡说》', '《量子废话》',
    '《膜宇宙学杂志》', '《弦论早餐版》', '《凝聚态胡说》',
    '《粒子物理扯淡》', '《天体物理废话》', '《宇宙学胡说》'
]

# 英文期刊
ENGLISH_JOURNALS = [
    'Journal of Nonsense Physics', 'Chaos and Disorder Letters',
    'Quantum Breakfast Dynamics', 'Applied Bullshit Research',
    'Theoretical Nonsense Review', 'Experimental Gibberish',
    'Modern Pseudoscience Today', 'Fictional Physics Letters',
    'Abstract Research Express', 'Nonlinear Bullshit Journal',
    'Statistical Nonsense', 'Quantum Gibberish',
    'Brane Cosmology Journal', 'String Theory Breakfast Edition',
    'Condensed Matter Nonsense', 'Particle Physics Bullshit',
    'Astrophysical Gibberish', 'Cosmological Nonsense Review',
    'Journal of Applied Chaos', 'International Journal of Nonsense'
]

# 研究方法
RESEARCH_METHODS = [
    '蒙特卡洛方法', '分子动力学模拟', '密度泛函理论', '第一性原理计算',
    '有限元分析', '边界元方法', '傅里叶变换', '小波分析',
    '主成分分析', '独立成分分析', '支持向量机', '随机森林',
    '神经网络', '深度学习', '强化学习', '贝叶斯推断',
    '马尔可夫链蒙特卡洛', '变分推断', '期望最大化', '梯度下降',
    '牛顿法', '拟牛顿法', '共轭梯度法', '线性规划',
    '非线性规划', '动态规划', '遗传算法', '粒子群优化',
    '模拟退火', '禁忌搜索', '蚁群算法', '免疫算法'
]

# 物理量/学术术语
PHYSICAL_QUANTITIES = [
    '熵', '焓', '自由能', '配分函数', '态密度', '费米能级',
    '玻色分布', '费米分布', '麦克斯韦分布', '玻尔兹曼分布',
    '关联函数', '结构因子', '响应函数', '极化率', '磁化率',
    '介电常数', '电导率', '热导率', '扩散系数', '黏滞系数',
    '表面张力', '接触角', '浸润性', '吸附能', '脱附能',
    '活化能', '反应速率', '半衰期', '寿命', '弛豫时间',
    '相干时间', '退相干时间', '纠缠熵', '保真度'
]

# 章节结构
CHAPTER_STRUCTURES = {
    'journal': [
        {
            'title': '引言',
            'sections': ['研究背景与意义', '研究现状与问题', '本文主要贡献', '论文组织结构']
        },
        {
            'title': '相关理论与文献综述',
            'sections': ['经典理论框架', '最新研究进展', '存在的问题与争议', '本文研究切入点']
        },
        {
            'title': '理论模型与方法',
            'sections': ['基本假设与定义', '理论模型构建', '数学推导', '数值方法']
        },
        {
            'title': '实验设计与方法',
            'sections': ['实验材料与设备', '实验方案设计', '实验流程', '数据处理方法']
        },
        {
            'title': '结果与分析',
            'sections': ['主要实验结果', '参数影响分析', '对比验证实验', '机制讨论']
        },
        {
            'title': '讨论',
            'sections': ['结果的理论意义', '与现有研究的对比', '研究局限性', '未来研究方向']
        },
        {
            'title': '结论',
            'sections': ['主要结论', '研究创新点', '应用前景展望']
        }
    ],
    'master': [
        {
            'title': '绪论',
            'sections': ['研究背景与意义', '国内外研究现状', '研究内容与方法', '论文组织结构']
        },
        {
            'title': '文献综述',
            'sections': ['相关理论基础', '国内外研究进展', '研究述评', '本文研究方向']
        },
        {
            'title': '理论基础',
            'sections': ['基本概念', '核心理论', '数学工具', '相关方法']
        },
        {
            'title': '研究方法',
            'sections': ['总体研究方案', '关键技术', '实验设计', '数据采集与处理']
        },
        {
            'title': '实验结果与分析',
            'sections': ['实验结果展示', '结果分析与讨论', '参数敏感性分析', '对比实验']
        },
        {
            'title': '综合讨论',
            'sections': ['理论解释', '实践意义', '研究局限', '未来展望']
        },
        {
            'title': '结论与展望',
            'sections': ['主要结论', '研究创新点', '不足与展望', '后续工作建议']
        }
    ],
    'doctor': [
        {
            'title': '绪论',
            'sections': ['研究背景', '研究意义', '问题提出', '研究目标与内容', '论文组织结构']
        },
        {
            'title': '文献综述',
            'sections': ['国外研究现状', '国内研究现状', '研究述评', '存在的问题', '本文研究方向']
        },
        {
            'title': '理论基础',
            'sections': ['基本概念', '经典理论', '现代发展', '数学工具']
        },
        {
            'title': '理论模型构建',
            'sections': ['基本假设', '模型建立', '理论推导', '数值方法']
        },
        {
            'title': '实验一：基础验证',
            'sections': ['实验目的', '实验设计', '实验过程', '结果与分析']
        },
        {
            'title': '实验二：参数研究',
            'sections': ['实验目的', '实验设计', '实验过程', '结果与分析']
        },
        {
            'title': '实验三：机制探索',
            'sections': ['实验目的', '实验设计', '实验过程', '结果与分析']
        },
        {
            'title': '综合分析与讨论',
            'sections': ['结果综合分析', '理论解释', '与现有研究对比', '研究贡献', '局限性分析']
        },
        {
            'title': '结论与展望',
            'sections': ['主要结论', '研究创新点', '理论意义', '应用价值', '未来研究方向']
        }
    ],
    'conference': [
        {
            'title': '引言',
            'sections': ['研究背景', '研究问题', '本文贡献']
        },
        {
            'title': '相关工作',
            'sections': ['已有研究', '存在问题', '本文方法']
        },
        {
            'title': '方法',
            'sections': ['总体架构', '核心算法', '实现细节']
        },
        {
            'title': '实验',
            'sections': ['实验设置', '实验结果', '对比分析']
        },
        {
            'title': '结论',
            'sections': ['总结', '未来工作']
        }
    ]
}

# 公式模板
FORMULA_TEMPLATES = [
    'E = mc²',
    'S = k_B ln Ω',
    'Δx · Δp ≥ ħ/2',
    'iħ ∂|ψ⟩/∂t = Ĥ|ψ⟩',
    'Z = Σ e^(-E_i/kT)',
    'F = -∇U',
    '∇·E = ρ/ε₀',
    '∇×B = μ₀J + μ₀ε₀ ∂E/∂t',
    'dS/dt ≥ 0',
    'λ = h/p',
    'T = 2π√(l/g)',
    'P = ρgh'
]

# 致谢人物
THANKS_PEOPLE = [
    '我的导师', '实验室的同学们', '楼下卖奶茶的阿姨', '外卖小哥',
    '我的猫', '我的狗', '我的室友', '我的父母',
    '张三', '李四', '王五', '赵六',
    '不知名的网友', 'B站UP主', '某游戏主播', '我的高中老师',
    '小区保安', '快递员', '食堂阿姨', '宿管大爷'
]

# 致谢物品
THANKS_THINGS = [
    '咖啡', '奶茶', '泡面', '螺蛳粉', '煎饼果子',
    '烧烤', '火锅', '可乐', '薯片', '巧克力',
    '耳机', '键盘', '显示器', '鼠标', '椅子',
    '空调', 'WiFi', '充电宝', '外卖', '夜宵'
]

# 胡说八道彩蛋内容
# 英文学术名词
ENGLISH_TERMS = [
    'Quantum Coherence', 'Topological Order', 'Emergent Phenomena',
    'Non-equilibrium Dynamics', 'Self-organized Criticality',
    'Fractal Dimension', 'Strange Attractor', 'Lyapunov Exponent',
    'Renormalization Group', 'Effective Field Theory',
    'Path Integral', 'Feynman Diagram', 'Density Functional Theory',
    'Molecular Dynamics', 'Monte Carlo Simulation',
    'Brownian Motion', 'Stochastic Process', 'Markov Chain',
    'Bayesian Inference', 'Maximum Entropy',
    'Phase Transition', 'Critical Exponent', 'Scaling Law',
    'Universality Class', 'Symmetry Breaking',
    'Goldstone Mode', 'Higgs Mechanism', 'Gauge Theory',
    'String Theory', 'M-theory', 'Brane Cosmology',
    'Dark Matter', 'Dark Energy', 'Cosmological Constant',
    'Multiverse Hypothesis', 'Anthropic Principle',
    'Quantum Gravity', 'Loop Quantum Gravity',
    'Holographic Principle', 'AdS/CFT Correspondence',
    'Quantum Entanglement', 'Quantum Teleportation',
    'Quantum Computing', 'Quantum Supremacy',
    'Neural Network', 'Deep Learning', 'Backpropagation',
    'Reinforcement Learning', 'Generative Adversarial Network',
    'Transformer Architecture', 'Attention Mechanism',
    'Emergent Behavior', 'Collective Intelligence',
    'Complex Adaptive System', 'Edge of Chaos',
    'Quantum Tunneling', 'Quantum Superposition', 'Wave-particle Duality',
    'Uncertainty Principle', 'Schrödinger Equation', 'Dirac Equation',
    'Maxwell Equations', 'Navier-Stokes Equations', 'Euler-Lagrange Equation',
    'Hamiltonian Formalism', 'Lagrangian Mechanics',
    'Statistical Mechanics', 'Thermodynamic Limit', 'Ergodic Hypothesis',
    'Poincaré Recurrence',
    'Chaos Theory', 'Bifurcation Theory', 'Catastrophe Theory',
    'Synergetics', 'Dissipative Structure', 'Self-organization',
    'Autopoiesis', 'Structural Coupling',
    'Observational Equivalence', 'Underdetermination', 'Duhem-Quine Thesis',
    'Paradigm Shift', 'Scientific Revolution', 'Normal Science',
    'Research Programme', 'Hard Core', 'Protective Belt',
    'Progressive Problemshift', 'Degenerating Problemshift',
    'Falsifiability', 'Verificationism', 'Confirmation Holism',
    'Theory-ladenness', 'Incommensurability', 'Conceptual Scheme',
    'Ontological Relativity', 'Indeterminacy of Translation', 'Gavagai',
    'Quantum Darwinism', 'Decoherence', 'Many-worlds Interpretation',
    'Copenhagen Interpretation', 'Pilot-wave Theory', 'Consistent Histories',
    'Quantum Bayesianism', 'Relational Quantum Mechanics',
    'Transactional Interpretation', 'Ensemble Interpretation',
    'Objective Collapse', 'Spontaneous Localization',
    'Penrose Interpretation', 'Orch-OR',
    'Quantum Mind', 'Hard Problem of Consciousness',
    'Philosophical Zombie', 'Qualia', 'Epiphenomenalism',
    'Panpsychism', 'Integrated Information Theory',
    'Global Workspace Theory', 'Multiple Drafts Model',
    'Chinese Room Argument', 'Turing Test', "Searle's Chinese Room",
    'Functionalism', 'Type Physicalism', 'Token Physicalism',
    'Anomalous Monism', 'Supervenience', 'Emergentism',
    'Reductionism', 'Holism', 'Methodological Individualism',
    'Methodological Holism', 'Microreduction', 'Macroreduction',
    'Intertheoretic Reduction', 'Bridge Laws', 'Special Sciences',
    'Multiple Realizability', 'Downward Causation',
    'Top-down Causation', 'Bottom-up Causation',
    'Circular Causality', 'Reciprocal Causation',
    'Mutual Manipulability', 'Causal Exclusion',
    'Overdetermination', 'Parallelism', 'Occasionalism',
    'Pre-established Harmony', 'Psychophysical Parallelism',
    'Mind-body Problem', 'Easy Problems', 'The Explanatory Gap',
    'Neural Correlates of Consciousness', 'NCC',
    'Global Neuronal Workspace', 'Local Recurrent Processing',
    'First-order Theories', 'Higher-order Theories',
    'Higher-order Thought', 'Higher-order Perception',
    'Self-representational Theories', 'Same-order Theories',
    'Intrinsic Theories', 'Quantum Theories of Consciousness',
    'Orchestrated Objective Reduction', 'Penrose-Hameroff',
    'Microtubules', 'Quantum Coherence in the Brain',
    'Warm Quantum Brain'
]

# 混沌学术词汇
CHAOS_WORDS = [
    '超验性', '存在论', '认识论', '本体论', '范式转换',
    '解构主义', '建构主义', '实证主义', '证伪主义',
    '现象学', '诠释学', '结构主义', '后结构主义',
    '话语实践', '知识考古', '谱系学', '异质性',
    '延异', '播散', '踪迹', '补充',
    '镜像阶段', '想象界', '象征界', '实在界',
    '力比多', '无意识', '潜意识', '集体无意识',
    '原型', '阴影', '人格面具', '阿尼玛',
    '能指', '所指', '漂浮的能指', '符号秩序',
    '规训', '惩戒', '生命权力', '治理术',
    '全景敞视主义', '异托邦', '拟像', '超真实',
    '内爆', '外爆', '媒介即讯息', '地球村',
    '赛博空间', '虚拟现实', '增强现实', '元宇宙',
    '后人类主义', '超人类主义', '技术奇点',
    '涌现性', '自组织', '耗散结构', '协同学',
    '突变论', '分形几何', '混沌理论', '复杂系统',
    '蝴蝶效应', '初值敏感性', '奇怪吸引子',
    '分岔', '倍周期', '准周期', '混沌边缘',
    '量子隧穿', '量子叠加', '量子纠缠',
    '薛定谔的猫', '观察者效应', '测量问题',
    '多世界诠释', '退相干', '量子芝诺效应',
    '本体论承诺', '认识论断裂', '知识型', '知识考古学',
    '权力/知识', '生物权力', '生命政治', '治理性',
    '国家理性', '警察国家', '规训社会', '控制社会',
    '全景敞视', '环形监狱', '边沁的全景监狱',
    '异质空间', '乌托邦', '反乌托邦', '敌托邦',
    '恶托邦', '无何有之乡', '乌有之乡', '世外桃源',
    '香格里拉', '伊甸园', '失乐园', '复乐园',
    '巴别塔', '诺亚方舟', '最后的审判', '启示录',
    '天启', '末日审判', '千禧年', '末世论', '终末论',
    '历史的终结', '最后的人', '超人', '权力意志',
    '永恒轮回', '重估一切价值', '上帝已死', '虚无主义',
    '消极虚无主义', '积极虚无主义', '最高价值的自行贬黜',
    '柏拉图主义', '反柏拉图主义', '形而上学', '存在者整体',
    '存在之真理', '存在的遗忘', '此在', '在世之在',
    '与他人共在', '向死而生', '畏', '烦', '沉沦',
    '被抛', '筹划', '本真状态', '非本真状态',
    '常人', '他们', '闲言', '好奇', '两可',
    '沉沦态', '畏之无', '无', '虚无',
    '存在与时间', '存在与虚无', '存在与事件', '事件',
    '增补', '替补', '补充逻辑', '在场形而上学',
    '逻各斯中心主义', '语音中心主义', '男性中心主义',
    '菲勒斯中心主义', '本质主义', '反本质主义',
    '基础主义', '反基础主义', '相对主义', '绝对主义',
    '普遍主义', '特殊主义', '个体主义', '整体主义',
    '原子主义', '单子论', '前定和谐', '充足理由律',
    '矛盾律', '同一律', '排中律', '二值原则',
    '多值逻辑', '模糊逻辑', '模态逻辑', '道义逻辑',
    '时态逻辑', '认知逻辑', '直觉主义逻辑', '相干逻辑',
    '次协调逻辑', '弗协调逻辑', '双面真理论', '爆炸原理',
    '归谬法', '反证法', '辩证法', '正题', '反题', '合题',
    '否定之否定', '扬弃', '异化', '对象化', '外化', '内化',
    '中介', '直接性', '间接性', '具体普遍性', '抽象普遍性',
    '具体同一性', '差异辩证法', '同一辩证法', '矛盾辩证法',
    '主奴辩证法', '承认的辩证法', '精神现象学', '绝对精神',
    '绝对理念', '绝对知识', '实体即主体',
    '密涅瓦的猫头鹰', '理性的狡计', '理性的机巧', '历史的狡计',
    '世界历史', '世界历史个人', '历史理性', '历史决定论',
    '历史目的论', '历史进步论', '历史循环论', '历史终结论',
    '后历史', '后现代', '后现代性', '后现代主义',
    '后殖民主义', '后马克思主义', '后哲学', '后形而上学',
    '后神学', '后真相', '后事实', '后真理', '后真相时代',
    '后真相政治', '后真相社会', '另类事实', '假新闻',
    '虚假信息', '信息茧房', '回声室效应', '过滤气泡',
    '算法偏见', '数据主义', '数据化', '量化自我',
    '监控资本主义', '监视资本主义', '平台资本主义',
    '认知资本主义', '非物质劳动', '数字劳动', '情感劳动',
    '关系劳动', '生命劳动', '活劳动', '死劳动',
    '可变资本', '不变资本', '剩余价值', '绝对剩余价值',
    '相对剩余价值', '超额剩余价值', '剩余价值率', '利润率',
    '平均利润率', '利润率下降趋势', '资本有机构成',
    '资本积累', '资本集中', '资本积聚', '资本循环',
    '资本周转', '固定资本', '流动资本', '货币资本',
    '生产资本', '商品资本', '产业资本', '商业资本',
    '借贷资本', '生息资本', '虚拟资本', '金融资本',
    '金融化', '新自由主义', '新保守主义', '新帝国主义',
    '新殖民主义', '新种族主义', '新性别主义', '新异性恋主义',
    '新父权制', '新资本主义', '晚期资本主义', '跨国资本主义',
    '全球资本主义', '帝国', '诸众', '生命政治生产',
    '非物质劳动霸权', '数字资本主义', '算法治理', '数据治理',
    '智能治理', '技术治理', '国家治理', '社会治理',
    '基层治理', '城市治理', '全球治理', '区域治理',
    '环境治理', '生态治理', '风险治理', '危机治理',
    '应急治理', '常态化治理', '运动式治理', '项目制治理',
    '网格化治理', '精细化治理', '精准治理', '智慧治理',
    '数字治理', '电子政务', '数字政府', '智慧政府',
    '服务型政府', '法治政府', '有限政府', '有效政府',
    '责任政府', '透明政府', '廉洁政府', '高效政府',
    '创新政府', '学习型政府', '创新型政府', '企业家政府',
    '企业化政府', '市场化政府', '新公共管理', '新公共服务',
    '新公共行政', '公共治理', '公共管理', '公共行政',
    '公共政策', '公共利益', '公共价值', '公共理性',
    '公共领域', '公共空间', '公共产品', '公共服务',
    '公共物品', '公共选择', '公共选择理论', '理性选择',
    '理性选择理论', '集体行动', '集体行动的逻辑', '搭便车',
    '公地悲剧', '囚徒困境', '零和博弈', '非零和博弈',
    '正和博弈', '负和博弈', '纳什均衡', '帕累托最优',
    '帕累托改进', '卡尔多-希克斯效率', '潜在帕累托改进',
    '社会选择', '社会选择理论', '阿罗不可能定理',
    '多数票规则', '孔多塞悖论', '投票悖论', '中位投票人定理',
    '唐斯模型', '理性无知', '理性冷漠', '理性弃权',
    '政治冷漠', '政治参与', '政治效能感', '政治信任',
    '政治合法性', '政治正当性', '政治权威', '政治权力',
    '政治权利', '政治自由', '政治平等', '政治正义',
    '政治民主', '政治发展', '政治现代化', '政治变迁',
    '政治转型', '政治转轨', '政治过渡', '政治稳定',
    '政治秩序', '政治制度化', '政治衰败', '政治腐化',
    '政治腐败', '政治俘获', '政治寻租', '政治租金',
    '政治市场', '政治企业家', '政治供给', '政治需求',
    '政治均衡', '政治周期', '政治商业周期', '机会主义',
    '道德风险', '逆向选择', '信号传递', '信号甄别',
    '委托-代理问题', '代理成本', '交易成本', '交易费用',
    '科斯定理', '产权理论', '契约理论', '不完全契约',
    '关系契约', '隐性契约', '显性契约', '自我执行契约',
    '声誉机制', '信任机制', '社会资本', '社会网络',
    '嵌入性', '弱关系', '强关系', '结构洞', '社会资源',
    '社会分层', '社会流动', '社会排斥', '社会整合',
    '社会团结', '社会凝聚力', '社会共识', '社会冲突',
    '社会矛盾', '社会问题', '社会风险', '社会危机',
    '社会动荡', '社会失序', '社会失范', '越轨行为',
    '偏差行为', '社会控制', '社会规范', '社会习俗',
    '社会道德', '社会伦理', '社会价值', '社会文化',
    '社会结构', '社会制度', '社会组织', '社会系统',
    '社会功能', '社会角色', '社会互动', '社会交往',
    '社会关系', '社会支持', '社会救助', '社会保障',
    '社会福利', '社会服务', '社会工作', '社会政策',
    '社会行政', '社会管理', '社会建设', '社会发展',
    '社会进步', '社会变迁', '社会转型', '社会转轨',
    '社会过渡', '社会现代化', '社会城市化', '社会工业化',
    '社会信息化', '社会数字化', '社会智能化', '社会网络化',
    '社会平台化', '社会算法化', '社会数据化', '社会量化',
    '社会可计算化', '社会可预测性', '社会可控制性',
    '社会可治理性', '社会工程', '社会设计', '社会规划',
    '社会实验', '社会试点', '社会示范', '社会创新',
    '社会创业', '社会企业', '社会团体', '社会中介',
    '社会第三部门', '社会非营利组织', '社会非政府组织',
    '社会公民社会', '社会市民社会', '社会公共领域',
    '社会公共空间', '社会公共舆论', '社会公共意见',
    '社会公共理性', '社会公共价值', '社会公共利益',
    '社会公共产品', '社会公共服务', '社会公共物品',
    '社会公共政策', '社会公共管理', '社会公共行政',
    '社会公共治理', '社会公共事务', '社会公共问题',
    '社会公共危机', '社会公共风险', '社会公共安全',
    '社会公共卫生', '社会公共教育', '社会公共文化',
    '社会公共交通', '社会公共设施', '社会公共资源',
    '社会公共环境', '社会公共生态', '社会公共福祉',
    '社会公共善', '社会公共正义', '社会公共公平',
    '社会公共平等', '社会公共自由', '社会公共民主',
    '社会公共法治', '社会公共秩序', '社会公共稳定',
    '社会公共和谐', '社会公共幸福', '社会公共美好',
    '社会公共理想', '社会公共愿景', '社会公共目标',
    '社会公共使命', '社会公共责任', '社会公共担当'
]

# 胡说八道彩蛋内容
# 胡说八道彩蛋内容
CRAZY_CONTENTS = [
    '实验数据会随星期几变化',
    '实验室的猫进来后结果就不一样了',
    '下雨天和晴天结果差异很大',
    '实验员心情好的时候数据更漂亮',
    '金鱼的游动与实验结果存在神秘关联',
    '植物的生长速度影响实验精度',
    '天气变化导致数据漂移',
    '股票涨跌与实验结果呈正相关',
    '播放古典音乐时稳定性显著提升',
    '播放摇滚乐时结果明显下降',
    '听相声时实验产率最高',
    '白噪音对信噪比有改善作用',
    '周一的实验结果总是特别差',
    '周五的数据质量明显下降',
    '满月时会观察到奇异现象',
    '新月时实验重复性最好',
    '每次导师来看实验就会出问题',
    '熬夜做的实验结果反而更好',
    '喝完奶茶后做实验准确率会提升',
    '设备会在最关键的时候出故障',
    '对照组的数据总是比实验组好看',
    '重复实验的结果永远不会完全一样',
    '论文投稿后才会发现数据有问题',
    '实验记录本会莫名其妙地消失',
    '备份的文件永远找不到',
    '以为是bug，结果是feature',
    '以为是feature，结果是bug',
    '数据会在你不看的时候自己变化',
    '实验结果与星座运势高度相关',
    '黄历上写着"不宜做实验"那天数据一定不好',
    '实验室的风水会影响实验结果',
    '每次换个座位，结果就不一样了',
    '用蓝色笔记录的数据更准确',
    '外卖到了的那一刻，数据就出问题了',
    '只要你说"这次应该没问题了"，就一定会出问题',
    '墨菲定律在实验室里永远成立',
    '越重要的实验，越容易出问题',
    '越简单的步骤，越容易出错',
    '你最有把握的地方，往往就是问题所在',
    '凌晨三点的实验数据总是最神奇',
    '实验室的微波炉会干扰实验结果',
    '有人在背后看你做实验的时候，一定会出错',
    '只要你一拍照，设备就坏了',
    '理论上应该是这样的，但实际上不是',
    '实际上是这样的，但理论上解释不通',
    '这个现象只有在你老板面前才会出现',
    '重复三次都一样的结果，一定是哪里错了',
    '不一样的结果，说明实验没做好',
    '反正最后都能凑出一篇论文来',
    '实验结果会随着实验员的发型变化',
    '穿红色衣服做实验数据会更喜庆',
    '穿黑色衣服做实验数据会更严肃',
    '穿白色衣服做实验数据会更纯净',
    '穿拖鞋做实验数据会更随意',
    '穿正装做实验数据会更正式',
    '实验前洗手数据会更干净',
    '实验前不洗手数据会更有味道',
    '实验前祈祷数据会更神圣',
    '实验前烧香数据会更有灵性',
    '实验前算命数据会更有预见性',
    '实验前看黄历数据会更符合天时',
    '实验前看星座数据会更符合运势',
    '实验前看塔罗牌数据会更神秘',
    '实验前抛硬币数据会更随机',
    '实验前掷骰子数据会更不确定',
    '实验前抽签数据会更有宿命感',
    '实验结果和实验员的智商成反比',
    '实验结果和实验员的颜值成正比',
    '实验结果和实验员的身高无关',
    '实验结果和实验员的体重负相关',
    '实验结果和实验员的睡眠质量正相关',
    '实验结果和实验员的饮食质量正相关',
    '实验结果和实验员的感情状态密切相关',
    '实验员谈恋爱的时候数据会更甜蜜',
    '实验员失恋的时候数据会更苦涩',
    '实验员单身的时候数据会更独立',
    '实验员追星的时候数据会更疯狂',
    '实验员追剧的时候数据会更有剧情',
    '实验员打游戏的时候数据会更有策略',
    '实验员刷短视频的时候数据会更碎片化',
    '实验员看直播的时候数据会更实时',
    '实验结果会受到月球引力的影响',
    '实验结果会受到太阳黑子的影响',
    '实验结果会受到宇宙射线的影响',
    '实验结果会受到中微子振荡的影响',
    '实验结果会受到引力波的影响',
    '实验结果会受到暗物质的影响',
    '实验结果会受到暗能量的影响',
    '实验结果会受到平行宇宙的干扰',
    '实验结果会受到时间旅行的影响',
    '实验结果会受到外星人的干预',
    '实验结果会受到未来人类的警告',
    '实验结果会受到高维生物的观察',
    '实验结果会受到量子涨落的影响',
    '实验结果会受到真空极化的影响',
    '实验结果会受到希格斯场的影响',
    '实验结果会受到夸克禁闭的影响',
    '实验结果会受到渐近自由的影响',
    '实验结果会受到规范对称性的影响',
    '实验结果会受到自发对称性破缺的影响',
    '实验结果会受到重整化群的影响',
    '实验结果会受到有效场论的影响',
    '实验结果会受到路径积分的影响',
    '实验结果会受到费曼图的影响',
    '实验结果会受到密度泛函理论的影响',
    '实验结果会受到分子动力学的影响',
    '实验结果会受到蒙特卡洛模拟的影响',
    '实验结果会受到布朗运动的影响',
    '实验结果会受到随机过程的影响',
    '实验结果会受到马尔可夫链的影响',
    '实验结果会受到贝叶斯推断的影响',
    '实验结果会受到最大熵原理的影响',
    '实验结果会受到相变的影响',
    '实验结果会受到临界指数的影响',
    '实验结果会受到标度律的影响',
    '实验结果会受到普适类的影响',
    '实验结果会受到戈德斯通模的影响',
    '实验结果会受到希格斯机制的影响',
    '实验结果会受到规范理论的影响',
    '实验结果会受到弦理论的影响',
    '实验结果会受到M理论的影响',
    '实验结果会受到膜宇宙学的影响',
    '实验结果会受到宇宙学常数的影响',
    '实验结果会受到多重宇宙假说的影响',
    '实验结果会受到人择原理的影响',
    '实验结果会受到量子引力的影响',
    '实验结果会受到圈量子引力的影响',
    '实验结果会受到全息原理的影响',
    '实验结果会受到AdS/CFT对应的影响',
    '实验结果会受到量子隐形传态的影响',
    '实验结果会受到量子计算的影响',
    '实验结果会受到量子霸权的影响',
    '实验结果会受到神经网络的影响',
    '实验结果会受到深度学习的影响',
    '实验结果会受到反向传播的影响',
    '实验结果会受到强化学习的影响',
    '实验结果会受到生成对抗网络的影响',
    '实验结果会受到Transformer架构的影响',
    '实验结果会受到注意力机制的影响',
    '实验结果会受到涌现行为的影响',
    '实验结果会受到集体智能的影响',
    '实验结果会受到复杂适应系统的影响',
    '实验结果会受到混沌边缘的影响'
]


# ============== 工具函数 ==============

def random_choice(arr):
    """从数组中随机选择一个元素"""
    return random.choice(arr)


def random_int(min_val, max_val):
    """生成随机整数"""
    return random.randint(min_val, max_val)


def random_float(min_val, max_val, decimals=2):
    """生成随机浮点数"""
    return round(random.uniform(min_val, max_val), decimals)


def random_exponential(base=10):
    """生成指数级随机数"""
    return int(base ** (random.random() * 3 + 1))


def process_keywords(keywords=None):
    """处理关键词，不足3个则随机补充

    Args:
        keywords: 用户输入的关键词列表

    Returns:
        包含3个关键词的列表
    """
    if keywords is None:
        keywords = []

    result = list(keywords)

    # 如果关键词不足3个，从词库中随机补充
    while len(result) < 3:
        pool_names = list(KEYWORD_POOLS.keys())
        pool_name = random_choice(pool_names)
        word = random_choice(KEYWORD_POOLS[pool_name])
        if word not in result:
            result.append(word)

    return result[:3]


# ============== 文档格式函数 ==============

def set_font(run, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN,
             size=None, bold=False):
    """设置字体

    Args:
        run: docx的Run对象
        cn_font: 中文字体
        en_font: 英文字体
        size: 字号（Pt）
        bold: 是否加粗
    """
    if size is None:
        size = Pt(BODY_FONT_SIZE)

    run.font.size = size
    run.font.bold = bold
    run.font.name = en_font

    r_pr = run._element.find(qn('w:rPr'))
    if r_pr is None:
        r_pr = OxmlElement('w:rPr')
        run._element.insert(0, r_pr)

    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.insert(0, r_fonts)

    r_fonts.set(qn('w:eastAsia'), cn_font)
    r_fonts.set(qn('w:ascii'), en_font)
    r_fonts.set(qn('w:hAnsi'), en_font)


def setup_page(section):
    """设置页面格式

    Args:
        section: docx的Section对象
    """
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(section, attr, Cm(MARGIN_CM))


def add_heading(doc, text, level=1):
    """添加标题

    Args:
        doc: docx的Document对象
        text: 标题文本
        level: 标题级别（1、2、3）

    Returns:
        段落对象
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    
    if level == 1:
        pf.space_before = Pt(18)
        pf.space_after = Pt(6)
        size = Pt(HEADING1_SIZE)
    elif level == 2:
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        size = Pt(HEADING2_SIZE)
    elif level == 3:
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        size = Pt(HEADING3_SIZE)
    else:
        pf.space_before = Pt(8)
        pf.space_after = Pt(4)
        size = Pt(BODY_FONT_SIZE)

    set_font(p.add_run(text), FONT_HEADING_CN, FONT_HEADING_EN, size, bold=True)
    return p


def add_body(doc, text, first_line_indent=True):
    """添加正文段落

    Args:
        doc: docx的Document对象
        text: 段落文本
        first_line_indent: 是否首行缩进

    Returns:
        段落对象
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    if first_line_indent:
        pf.first_line_indent = Cm(0.74)

    set_font(p.add_run(text))
    return p


def add_figure_caption(doc, text):
    """添加图题

    Args:
        doc: docx的Document对象
        text: 图题文本

    Returns:
        段落对象
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(12)
    set_font(p.add_run(text), size=Pt(CAPTION_SIZE), bold=True)
    return p


def add_table_caption(doc, text):
    """添加表题

    Args:
        doc: docx的Document对象
        text: 表题文本

    Returns:
        段落对象
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    set_font(p.add_run(text), size=Pt(CAPTION_SIZE), bold=True)
    return p


def add_formula(doc, text):
    """添加公式

    Args:
        doc: docx的Document对象
        text: 公式文本

    Returns:
        段落对象
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    set_font(p.add_run(text), en_font=FONT_FORMULA)
    return p


def add_reference(doc, text):
    """添加参考文献

    Args:
        doc: docx的Document对象
        text: 参考文献文本

    Returns:
        段落对象
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(3)
    pf.left_indent = Cm(0.74)
    pf.first_line_indent = Cm(-0.74)
    set_font(p.add_run(text), size=Pt(CAPTION_SIZE))
    return p


# ============== 内容生成函数 ==============

def generate_title(kw1, kw2, kw3):
    """生成论文标题

    Args:
        kw1: 关键词1
        kw2: 关键词2
        kw3: 关键词3

    Returns:
        论文标题
    """
    templates = [
        f'基于{kw1}的{kw2}研究：{kw3}视角',
        f'{kw2}的{kw3}效应：基于{kw1}的分析',
        f'从{kw1}看{kw2}的{kw3}特性',
        f'{kw1}视角下的{kw2}：{kw3}的作用',
        f'论{kw1}与{kw2}的{kw3}相关性',
        f'{kw1}诱导的{kw2}{kw3}机制研究',
        f'{kw2}中的{kw1}效应：{kw3}理论解释',
        f'基于{kw3}的{kw1}与{kw2}协变研究',
        f'{kw1}、{kw2}与{kw3}的交叉学科研究',
        f'{kw2}的{kw1}调控：{kw3}模型'
    ]
    return random_choice(templates)


def generate_abstract(kw1, kw2, kw3, level=5):
    """生成摘要

    Args:
        kw1: 关键词1
        kw2: 关键词2
        kw3: 关键词3
        level: 胡说八道程度（1-10）

    Returns:
        摘要文本
    """
    intensity = level / 10
    abstract_parts = []

    # 研究背景
    abstract_parts.append(
        f'针对{random_choice(["长期以来被忽视的", "学术界普遍关注的", "具有重要理论意义的", "充满争议的"])}'
        f'{kw2}问题，本研究突破传统{random_choice(["理论框架", "研究范式", "分析方法", "实验手段"])}的局限，'
        f'首次引入"{kw1}{random_choice(["诱导", "驱动", "调控", "介导"])}'
        f'{kw3}"模型。'
    )

    # 研究方法
    abstract_parts.append(
        f'通过对{random_int(10, 100)}组'
        f'{random_choice(["标准实验周期", "对照实验", "观测样本", "数值模拟"])}的'
        f'{random_choice(["系统观测", "深入分析", "精确测量", "定量研究"])}，'
        f'结合{random_choice(RESEARCH_METHODS)}、{random_choice(RESEARCH_METHODS)}'
        f'与{random_choice(RESEARCH_METHODS)}，'
    )

    # 研究发现
    abstract_parts.append(
        f'发现{kw1}与{kw2}之间存在显著的'
        f'{random_choice(["非线性", "非单调", "非局域", "非平衡"])}'
        f'{random_choice(["共振", "耦合", "关联", "协变"])}关系'
        f'（r = {random_float(-0.95, 0.95, 2)}, '
        f'p < {random_choice(["0.001", "0.0001", "1e-5", "1e-6"])}），'
    )

    # 机制解释
    abstract_parts.append(
        f'且该关系在{random_choice(["特定条件下", "临界温度附近", "强场作用下", "量子极限"])}被'
        f'{random_choice(["显著增强", "明显抑制", "完全改变", "量子化"])}。'
    )

    # 理论贡献
    if intensity > 0.3:
        abstract_parts.append(
            f'进一步，本研究构建了"{kw1}-{kw2}联合'
            f'{random_choice(["熵", "能谱", "相图", "态密度"])}"，'
            f'并引入{random_choice(["量子场论", "弦论", "膜宇宙模型", "拓扑场论"])}'
            f'与{random_choice(["分数统计", "非交换几何", "非对易时空", "扭量理论"])}，'
            f'证明{kw2}并未{random_choice(["消失", "改变", "跃迁", "隧穿"])}，'
            f'而是通过{random_choice(["质能转化", "维度紧致化", "量子退相干", "拓扑相变"])}机制'
            f'转化为{random_choice(["能量", "信息", "熵", "暗物质"])}。'
        )

    # 意外发现
    if intensity > 0.6:
        abstract_parts.append(
            f'此外，研究还意外发现'
            f'{random_choice(["实验员心情", "天气状况", "星座运势", "午餐吃什么"])}'
            f'与实验结果存在{random_int(12, 48)}小时的'
            f'{random_choice(["超日同步", "相位锁定", "因果关联", "量子纠缠"])}，'
            f'提示存在一种{random_choice(["跨尺度", "跨维度", "跨物种", "跨宇宙"])}的'
            f'{random_choice(["宏观量子效应", "非局域关联", "长程有序", "自组织临界"])}现象。'
        )

    # 应用价值
    abstract_parts.append(
        f'本发现为{random_choice(["减少", "优化", "调控", "预测"])}'
        f'{kw2}提供了基于{kw1}选择与{kw3}规避的切实干预方案，'
        f'并初步勾勒了统一解释'
        f'{random_choice(["日常神秘现象", "量子-经典过渡", "暗物质本质", "时间箭头"])}的'
        f'"大统一{random_choice(["早餐", "奶茶", "螺蛳粉", "煎饼果子"])}理论"。'
    )

    return ''.join(abstract_parts)


def generate_keywords(kw1, kw2, kw3):
    """生成关键词

    Args:
        kw1: 关键词1
        kw2: 关键词2
        kw3: 关键词3

    Returns:
        关键词列表
    """
    extra_keywords = [
        random_choice(PHYSICAL_QUANTITIES),
        random_choice(PHYSICAL_QUANTITIES),
        random_choice([
            '非平衡态', '临界现象', '自组织', '涌现行为',
            '分形结构', '混沌动力学', '量子相干', '拓扑保护'
        ])
    ]

    return [kw1, kw2, kw3] + extra_keywords


def generate_authors_and_depts(n_authors=None):
    """生成作者和机构

    Args:
        n_authors: 作者数量，None则随机

    Returns:
        (authors, departments) 元组
    """
    if n_authors is None:
        n_authors = random_int(2, 6)

    n_depts = random_int(2, 4)
    departments = []
    for _ in range(n_depts):
        departments.append(random_choice(UNIVERSITIES))

    authors = random.sample(AUTHOR_NAMES, min(n_authors, len(AUTHOR_NAMES)))

    return authors, departments


def generate_paragraph(kw1, kw2, kw3, section_title, chapter_title, level=5):
    """生成正文段落

    Args:
        kw1: 关键词1
        kw2: 关键词2
        kw3: 关键词3
        section_title: 小节标题
        chapter_title: 章节标题
        level: 胡说八道程度（1-10）

    Returns:
        段落文本
    """
    intensity = level / 10
    sentences = []
    used_templates = set()  # 记录已使用的模板类型，避免重复

    # 开头句模板生成函数
    def get_openings():
        return [
            # 类型1：重要性阐述
            f'{section_title}是{chapter_title}的'
            f'{random_choice(["核心内容", "重要组成", "关键部分", "理论基础", "方法支撑", "前沿领域", "热点问题", "难点所在"])}，'
            f'对理解{kw1}与{kw2}的关系具有'
            f'{random_choice(["重要", "深远", "关键", "决定性", "基础性", "开创性"])}意义。',
            f'在{chapter_title}的研究框架中，{section_title}始终占据着'
            f'{random_choice(["核心", "重要", "关键", "特殊", "独特"])}地位。',
            f'{section_title}作为连接{kw1}与{kw3}的桥梁，在本研究中扮演着'
            f'{random_choice(["关键", "重要", "核心", "特殊"])}角色。',
            f'深入研究{section_title}，对于揭示{kw2}的'
            f'{random_choice(["本质", "内在机制", "演化规律", "物理本质"])}具有重要的理论价值。',

            # 类型2：研究现状
            f'关于{kw2}的{section_title}，学术界一直存在'
            f'{random_choice(["广泛讨论", "诸多争议", "不同看法", "多种观点", "较大分歧", "激烈争论"])}。',
            f'{section_title}的研究可以追溯到{random_int(1900, 2000)}年代，'
            f'当时{random_choice(AUTHOR_NAMES)}等人首次提出了{kw1}的概念。',
            f'近{random_int(10, 30)}年来，{section_title}的研究取得了'
            f'{random_choice(["长足进展", "显著成果", "突破性进展", "重要发现"])}。',
            f'尽管{section_title}已经被研究了多年，但仍有许多'
            f'{random_choice(["关键问题", "未解之谜", "核心机制", "基本问题"])}有待澄清。',
            f'随着{random_choice(["实验技术", "理论方法", "计算能力", "观测手段"])}的进步，'
            f'{section_title}的研究进入了一个新的阶段。',

            # 类型3：引入过渡
            f'在深入讨论{section_title}之前，有必要先明确几个基本概念。',
            f'为了更好地理解后续内容，本节首先对{section_title}进行简要介绍。',
            f'本章将重点围绕{section_title}展开'
            f'{random_choice(["深入", "系统", "全面", "详细"])}讨论。',
            f'接下来，我们将从{random_choice(["理论", "实验", "数值", "分析"])}角度探讨{section_title}。',

            # 类型4：问题导向
            f'{kw2}的{section_title}一直是'
            f'{random_choice(["学术界", "工业界", "理论界", "实验界"])}关注的焦点问题。',
            f'如何准确{random_choice(["描述", "预测", "控制", "解释"])}'
            f'{kw1}对{kw2}的影响，是{section_title}的核心问题。',
            f'{section_title}的核心挑战在于如何处理{kw1}与{kw3}之间的'
            f'{random_choice(["耦合", "相互作用", "竞争", "协同"])}效应。',

            # 类型5：本文工作
            f'本节将重点介绍我们在{section_title}方面的'
            f'{random_choice(["最新进展", "研究成果", "主要工作", "创新点"])}。',
            f'针对{section_title}存在的问题，本文提出了一种新的'
            f'{random_choice(["理论模型", "实验方法", "分析框架", "计算方案"])}。',
            f'与现有研究不同，本文从{kw3}的角度重新审视了{section_title}。'
        ]

    # 理论阐述模板生成函数
    def get_theory_sentences():
        return [
            # 理论框架类
            f'根据{random_choice(["经典理论", "量子力学", "统计物理", "热力学第二定律", "朗道理论", "金兹堡-朗道方程", "重整化群理论", "标度理论", "对称性原理", "守恒定律"])}，'
            f'{kw1}的{random_choice(["演化", "涨落", "输运", "弛豫", "相变", "临界行为", "动力学", "热力学"])}过程可以用'
            f'{random_choice(["偏微分方程", "积分方程", "微分方程组", "变分原理", "哈密顿量", "拉格朗日量", "配分函数", "密度矩阵"])}来描述。',

            # 尺度分析类
            f'从{random_choice(["微观", "宏观", "介观", "纳观", "原子尺度", "分子尺度", "连续介质"])}尺度来看，'
            f'{kw2}的{random_choice(["结构", "性质", "行为", "响应", "演化", "稳定性", "相变"])}本质上是由'
            f'{kw1}的{random_choice(["集体行为", "统计分布", "量子态", "相互作用", "关联效应", "涨落", "相干性"])}决定的。',

            # 原理定律类
            f'按照{random_choice(["对称性原理", "守恒定律", "最小作用量原理", "熵增原理", "能量守恒", "动量守恒", "角动量守恒", "电荷守恒", "刘维尔定理", "细致平衡原理"])}，'
            f'系统会自发地趋向于'
            f'{random_choice(["能量最低", "熵最大", "自由能最小", "最概然", "平衡态", "稳态", "基态"])}的状态。',

            # 条件状态类
            f'在{random_choice(["平衡态", "非平衡态", "稳态", "瞬态", "临界状态", "亚稳态", "激发态", "基态"])}条件下，'
            f'{kw1}与{kw2}之间满足'
            f'{random_choice(["线性响应", "非线性耦合", "标度关系", "幂律分布", "指数关系", "对数关系", "饱和关系"])}。',

            # 模型构建类
            f'为了描述这一现象，我们建立了一个基于{kw1}和{kw2}的'
            f'{random_choice(["有效模型", "唯象模型", "微观模型", "连续模型", "离散模型", "平均场模型"])}。',
            f'该理论的核心假设是{kw3}可以作为'
            f'{random_choice(["序参量", "控制参数", "耦合常数", "特征尺度", "特征时间"])}来描述系统的行为。',

            # 数学描述类
            f'从数学上看，{kw1}的运动方程可以表示为一个'
            f'{random_choice(["非线性", "线性", "耦合", "耗散", "保守"])}系统，其解具有'
            f'{random_choice(["丰富的", "复杂的", "多样的", "奇异的"])}动力学行为。',
            f'通过引入{random_choice(["无量纲化", "标度变换", "坐标变换", "正则变换"])}，'
            f'可以将问题简化为{random_int(1, 5)}个独立参数的函数。',

            # 物理图像类
            f'物理上，这一现象可以理解为{kw1}与{kw2}之间通过{kw3}发生了'
            f'{random_choice(["能量交换", "动量传递", "信息传递", "相位耦合", "量子纠缠", "经典关联"])}。',
            f'直观地说，{kw2}的变化相当于在{kw1}的'
            f'{random_choice(["势能面", "能谱", "相空间", "参数空间"])}中引入了一个'
            f'{random_choice(["扰动", "调制", "势垒", "势阱"])}。'
        ]

    # 实验描述模板生成函数
    def get_experiment_sentences():
        return [
            # 实验设计类
            f'为了验证这一假设，我们设计了{random_int(3, 10)}组对照实验，'
            f'系统地改变了{random_choice(["温度", "压力", "浓度", "场强", "频率", "功率", "时间", "角度"])}等参数。',
            f'实验采用了{random_choice(["控制变量法", "正交设计", "响应面法", "田口方法", "全因子设计"])}，'
            f'以确保结果的{random_choice(["可靠性", "准确性", "可重复性", "统计显著性"])}。',

            # 设备技术类
            f'实验采用了{random_choice(["高精度", "高分辨率", "高灵敏度", "高速", "原位", "实时", "非接触式"])}测量设备，'
            f'采样频率达到{random_exponential()} Hz。',
            f'我们使用了{random_choice(["原子力显微镜", "扫描隧道显微镜", "透射电子显微镜", "X射线衍射仪", "拉曼光谱仪", "荧光光谱仪", "核磁共振仪"])}对样品进行了表征。',

            # 环境条件类
            f'所有实验均在{random_choice(["恒温恒湿", "超净", "真空", "低温", "高温", "高压", "强磁场"])}条件下进行，'
            f'每个条件重复测量{random_int(5, 20)}次以确保统计显著性。',
            f'为了排除环境干扰，实验装置放置在'
            f'{random_choice(["防震台", "屏蔽室", "恒温箱", "真空腔"])}中。',

            # 样品制备类
            f'实验样品采用{random_choice(["溶胶-凝胶法", "水热法", "气相沉积", "分子束外延", "电化学沉积", "机械合金化"])}方法制备，'
            f'样品尺寸约为{random_float(1, 1000, 2)} {random_choice(["nm", "μm", "mm", "cm"])}。',
            f'样品在测量前经过了{random_choice(["退火", "淬火", "抛光", "清洗", "镀膜"])}处理，'
            f'以确保表面质量。',

            # 数据处理类
            f'数据处理采用了{random_choice(RESEARCH_METHODS)}方法，'
            f'对原始数据进行了{random_choice(["滤波", "降噪", "归一化", "基线校正", "平滑", "拟合", "傅里叶变换", "小波分析"])}处理。',
            f'为了提高数据质量，我们采用了'
            f'{random_choice(["多次平均", "锁相放大", "相关检测", "差分测量"])}技术。',

            # 实验流程类
            f'实验过程中，首先{random_choice(["制备样品", "校准仪器", "设置参数", "建立基线"])}，'
            f'然后{random_choice(["施加激励", "改变条件", "开始测量", "采集数据"])}，'
            f'最后{random_choice(["分析结果", "处理数据", "验证结论", "重复实验"])}。',
            f'每个实验周期持续约{random_int(1, 24)}小时，期间'
            f'{random_choice(["实时监测", "定时采样", "连续记录", "间歇测量"])}各项参数。'
        ]

    # 结果描述模板生成函数
    def get_result_sentences():
        return [
            # 总体结论类
            f'实验结果表明，{kw1}对{kw2}具有'
            f'{random_choice(["显著的", "明显的", "微妙的", "出乎意料的", "复杂的", "非线性的", "多尺度的"])}影响。',
            f'我们的研究发现，{kw1}与{kw2}之间存在'
            f'{random_choice(["正相关", "负相关", "非线性相关", "非单调相关", "周期性相关"])}关系。',

            # 具体现象类
            f'我们观察到，当{random_choice(["温度", "压力", "浓度", "场强", "频率"])}达到'
            f'{random_float(0.1, 1000, 2)}时，{kw2}会发生'
            f'{random_choice(["突变", "相变", "共振", "击穿", "饱和", "转变", "分岔", "混沌"])}现象。',
            f'在{kw3}的作用下，{kw2}表现出'
            f'{random_choice(["反常行为", "奇异特性", "量子效应", "非单调变化", "滞后效应", "记忆效应"])}。',

            # 定量结果类
            f'定量分析显示，两者之间的相关系数高达{random_float(0.8, 0.999, 3)}，'
            f'表明存在{random_choice(["强相关", "显著相关", "高度显著相关", "极强相关", "统计显著相关"])}。',
            f'实验测得的{kw2}的{random_choice(["临界值", "特征值", "阈值", "饱和值", "共振频率"])}'
            f'约为{random_float(0.001, 1000, 3)}，与理论预测'
            f'{random_choice(["基本一致", "高度吻合", "存在偏差", "完全不同"])}。',

            # 对比验证类
            f'这一结果与{random_choice(["理论预测", "数值模拟", "前期研究", "直觉判断", "实验预期", "文献报道"])}'
            f'{random_choice(["高度一致", "基本吻合", "存在偏差", "完全不同", "部分一致", "趋势相同"])}。',
            f'与{random_choice(["传统方法", "现有技术", "前人工作", "其他方案"])}相比，'
            f'我们的方法具有{random_choice(["更高的精度", "更广的适用范围", "更低的计算成本", "更强的鲁棒性", "更好的稳定性"])}。',

            # 意外发现类
            f'令人意外的是，在{random_choice(["高温", "低温", "高压", "强场", "低频", "高频"])}条件下，'
            f'{kw2}表现出了{random_choice(["完全不同", "出乎意料", "反常", "奇异"])}的行为。',
            f'实验中还观察到了一些{random_choice(["有趣的", "奇怪的", "意外的", "反常的"])}现象，'
            f'例如{random_choice(["滞后回线", "记忆效应", "多稳态", "振荡行为"])}等。'
        ]

    # 机制讨论模板生成函数
    def get_mechanism_sentences():
        return [
            # 机制假说类
            f'关于这一现象的物理机制，我们认为可能与'
            f'{random_choice(["量子隧穿", "能级杂化", "界面态", "缺陷态", "声子散射", "电子-声子相互作用", "自旋-轨道耦合", "量子干涉", "多体效应", "拓扑保护"])}有关。',
            f'进一步分析表明，{kw3}在其中扮演了'
            f'{random_choice(["关键媒介", "催化剂", "抑制剂", "调控因子", "桥梁", "中介"])}的角色。',

            # 理论解释类
            f'这一效应可以用{random_choice(["朗道相变理论", "金兹堡-朗道方程", "重整化群", "标度理论", "平均场理论", "有效场论", "线性响应理论", "涨落-耗散定理"])}来解释。',
            f'从{random_choice(["热力学", "动力学", "统计力学", "量子力学", "经典力学"])}的角度来看，'
            f'这一现象的本质是{random_choice(["能量最小化", "熵最大化", "对称性破缺", "相变", "临界现象", "自组织"])}的结果。',

            # 微观机制类
            f'我们推测，其微观机制可能涉及'
            f'{random_choice(["电子-声子相互作用", "自旋-轨道耦合", "量子干涉", "多体效应", "缺陷散射", "界面散射", "能带结构", "态密度"])}。',
            f'在微观层面，{kw1}通过'
            f'{random_choice(["交换作用", "偶极相互作用", "范德华力", "氢键", "共价键", "静电相互作用"])}与{kw2}发生耦合。',

            # 能量角度类
            f'从能量角度分析，{kw2}的变化对应于系统'
            f'{random_choice(["自由能", "内能", "焓", "吉布斯自由能", "亥姆霍兹自由能"])}的'
            f'{random_choice(["降低", "升高", "重新分布", "极值"])}。',
            f'这一过程的{random_choice(["活化能", "能垒", "势阱深度", "特征能量"])}'
            f'约为{random_float(0.001, 100, 3)} eV，与'
            f'{random_choice(["实验测量值", "理论计算值", "文献报道值"])}基本一致。',

            # 动力学类
            f'动力学分析表明，{kw2}的{random_choice(["弛豫时间", "响应时间", "特征时间", "寿命"])}'
            f'与{kw1}的强度呈{random_choice(["指数关系", "幂律关系", "线性关系", "对数关系"])}。',
            f'这一效应的时间尺度约为{random_float(0.001, 1000, 3)} '
            f'{random_choice(["fs", "ps", "ns", "μs", "ms", "s"])}，属于'
            f'{random_choice(["超快", "快速", "中速", "慢速"])}过程。'
        ]

    # 引用文献模板生成函数
    def get_citation_sentences():
        return [
            # 某人研究表明类
            f'{random_choice(AUTHOR_NAMES)}等人[{random_int(1, 30)}]的研究表明，'
            f'{kw1}在{kw3}条件下会表现出'
            f'{random_choice(["反常行为", "奇异特性", "量子效应", "非单调变化", "临界行为", "非线性响应"])}。',
            f'{random_choice(AUTHOR_NAMES)}和{random_choice(AUTHOR_NAMES)}[{random_int(1, 30)}]通过'
            f'{random_choice(["实验", "理论", "数值模拟"])}研究了{kw2}的'
            f'{random_choice(["性质", "行为", "机制", "演化"])}。',

            # 期刊发表类
            f'根据{random_choice(JOURNALS)}上发表的最新研究[{random_int(1, 30)}]，'
            f'{kw2}的{random_choice(["临界温度", "转变点", "阈值", "特征尺度", "耦合强度"])}'
            f'约为{random_float(0.01, 1000, 3)}。',
            f'近期发表在{random_choice(JOURNALS)}上的一篇综述[{random_int(1, 30)}]'
            f'系统总结了{kw1}研究的最新进展。',

            # 证据支持类
            f'这一结论得到了大量实验证据的支持[{random_int(1, 30)}-{random_int(31, 50)}]。',
            f'许多研究小组都独立观察到了类似的现象[{random_int(1, 30)}, {random_int(1, 30)}, {random_int(1, 30)}]。',

            # 对比类
            f'与{random_choice(AUTHOR_NAMES)}等人的工作[{random_int(1, 30)}]相比，'
            f'本文的方法具有{random_choice(["更高的精度", "更广的适用范围", "更低的计算成本", "更强的鲁棒性", "更好的物理图像"])}。',
            f'不同于{random_choice(AUTHOR_NAMES)}等人提出的模型[{random_int(1, 30)}]，'
            f'我们的理论考虑了{kw3}的影响。',

            # 经典工作类
            f'早在{random_int(1950, 2000)}年，{random_choice(AUTHOR_NAMES)}就提出了著名的'
            f'{random_choice(["理论", "模型", "公式", "定律", "原理"])}[{random_int(1, 30)}]，'
            f'为后续研究奠定了基础。',
            f'经典的{random_choice(["理论", "模型", "实验"])}[{random_int(1, 30)}]认为'
            f'{kw2}只与{kw1}有关，但最新研究表明{kw3}也起着重要作用。',

            # 争议类
            f'然而，这一观点也受到了一些学者的质疑[{random_int(1, 30)}]，'
            f'他们认为{random_choice(["证据不足", "解释有误", "实验有缺陷", "理论不严谨"])}。',
            f'关于{kw2}的{random_choice(["起源", "机制", "本质", "应用"])}，'
            f'学术界仍存在不同看法[{random_int(1, 30)}, {random_int(1, 30)}]。'
        ]

    # 胡说八道彩蛋模板生成函数
    def get_crazy_sentences():
        return [
            # 时间/日期相关
            f'值得注意的是，实验过程中还观察到了一些'
            f'{random_choice(["有趣的", "奇怪的", "意外的", "神秘的", "难以解释的"])}现象，'
            f'比如{random_choice(CRAZY_CONTENTS)}，'
            f'其原因尚不清楚，有待进一步研究。',

            # 意外关联
            f'此外，我们还意外地发现'
            f'{random_choice(["金鱼的游动", "植物的生长", "天气变化", "股票涨跌", "网络流量", "社交媒体热度", "外卖订单量", "奶茶销量", "快递配送速度", "食堂排队长度"])}'
            f'与实验结果之间存在某种'
            f'{random_choice(["神秘的", "微妙的", "难以解释的", "统计学上显著的", "反直觉的", "非线性的", "量子化的"])}关联。',

            # 音乐/声音影响
            f'更令人惊讶的是，当我们播放'
            f'{random_choice(["古典音乐", "摇滚乐", "相声", "白噪音", "ASMR", "电子音乐", "爵士乐", "民谣", "重金属", "说唱"])}时，'
            f'{kw1}的{random_choice(["稳定性", "重复性", "信噪比", "产率", "精度", "效率"])}会'
            f'{random_choice(["显著提升", "明显下降", "发生变化", "变得不稳定", "出现周期性波动"])}。',

            # 特殊日子
            f'我们还注意到一个有趣的现象：'
            f'{random_choice(["周一", "周五", "满月", "新月", "节假日", "发薪日", "双十一", "618", "情人节", "圣诞节"])}的实验结果总是'
            f'{random_choice(["特别好", "特别差", "很奇怪", "很稳定", "波动很大"])}，'
            f'虽然我们不知道为什么。',

            # 人物影响
            f'有趣的是，当{random_choice(["导师", "实验室主任", "保洁阿姨", "外卖小哥", "楼下保安", "隔壁实验室的同学", "校长", "投资人"])}在场时，'
            f'实验结果会{random_choice(["明显变好", "明显变差", "变得稳定", "出现异常"])}，'
            f'这可能与{random_choice(["心理因素", "气场", "量子纠缠", "蝴蝶效应", "观察者效应"])}有关。',

            # 食物影响
            f'我们还发现，实验员{random_choice(["喝奶茶后", "吃火锅后", "吃螺蛳粉后", "喝咖啡后", "吃烧烤后", "没吃饭时", "吃饱后"])}做的实验，'
            f'数据质量{random_choice(["明显更高", "明显更低", "更稳定", "波动更大"])}。',

            # 设备玄学
            f'奇怪的是，{random_choice(["新买的仪器", "老仪器", "左边那台", "右边那台", "进口的", "国产的", "借来的"])}测出来的结果总是'
            f'{random_choice(["更准确", "更不准", "偏大", "偏小", "更稳定"])}，'
            f'尽管它们是同一型号。',

            # 网络玄学
            f'实验数据似乎与{random_choice(["网速", "WiFi信号强度", "手机电量", "电脑配置", "操作系统版本", "键盘颜色", "鼠标品牌"])}有某种神秘的联系。',

            # 墨菲定律类
            f'值得一提的是，每当我们{random_choice(["觉得这次应该没问题了", "准备向导师汇报", "要拍照记录", "准备收工"])}的时候，'
            f'实验就会{random_choice(["出问题", "失败", "数据异常", "设备故障"])}，完美验证了墨菲定律。',

            # 玄学总结类
            f'总的来说，{random_choice(["做实验就像开盲盒", "科研就是玄学", "数据全靠缘分", "结果全看运气"])}，'
            f'只要{random_choice(["样本量足够大", "重复次数足够多", "脸皮足够厚", "导师足够好"])}，总能凑出一篇论文。'
        ]

    # 混沌哲学类模板生成函数
    def get_chaos_sentences():
        return [
            # 存在论/本体论类
            f'从{random_choice(["存在论", "本体论", "认识论", "形而上学"])}的视角来看，'
            f'{kw1}的本质并非{random_choice(["实体", "属性", "关系", "过程"])}，'
            f'而是一种{random_choice(["涌现性", "建构性", "解构性", "超验性"])}的{random_choice(["存在方式", "话语实践", "符号秩序", "权力效应"])}。',

            # 范式转换类
            f'本研究实际上构成了一次{random_choice(["范式转换", "认识论断裂", "话语转向", "方法论革命"])}，'
            f'它挑战了传统的{random_choice(["主客二分", "本质主义", "基础主义", "逻各斯中心主义"])}，'
            f'开启了{kw2}研究的{random_choice(["新维度", "新视野", "新范式", "新问题域"])}。',

            # 解构/建构类
            f'通过对{kw1}的{random_choice(["解构", "解构主义分析", "谱系学考察", "知识考古"])}，'
            f'我们揭示了其背后隐藏的{random_choice(["权力结构", "意识形态", "话语策略", "文化密码"])}，'
            f'以及{kw3}在其中扮演的{random_choice(["共谋角色", "颠覆角色", "中介角色", "镜像角色"])}。',

            # 现象学/诠释学类
            f'从{random_choice(["现象学", "诠释学", "存在主义", "解释学"])}的立场出发，'
            f'{kw2}并非{random_choice(["客观事实", "纯粹现象", "中立数据", "自然对象"])}，'
            f'而是{random_choice(["意义的建构", "理解的产物", "存在的展开", "诠释的循环"])}。',

            # 符号/结构类
            f'在{random_choice(["符号秩序", "象征界", "话语体系", "文化编码"])}中，'
            f'{kw1}作为一个{random_choice(["能指", "符号", "隐喻", "转喻"])}，'
            f'其{random_choice(["意义", "所指", "指称", "内涵"])}始终处于{random_choice(["延异", "播散", "滑动", "漂移"])}状态。',

            # 精神分析类
            f'如果借用{random_choice(["精神分析", "拉康主义", "荣格心理学", "弗洛伊德主义"])}的概念框架，'
            f'我们可以说{kw2}是{random_choice(["无意识的", "被压抑的", "投射性的", "镜像性的"])}，'
            f'它反映了{kw3}的{random_choice(["阴影", "力比多", "集体无意识", "人格面具"])}。',

            # 福柯/权力类
            f'从{random_choice(["福柯", "生命政治", "治理术", "规训社会"])}的角度审视，'
            f'{kw1}的生产与传播实际上是一种{random_choice(["权力技术", "治理策略", "规训手段", "生命权力运作"])}，'
            f'它通过{random_choice(["分类", "标准化", "规范化", "主体化"])}来实现对{kw2}的管控。',

            # 鲍德里亚/拟像类
            f'在{random_choice(["拟像时代", "超真实社会", "媒介社会", "消费社会"])}中，'
            f'{kw2}已经不再是{random_choice(["对现实的反映", "真实的表象", "客观的存在", "自然的产物"])}，'
            f'而是一种{random_choice(["超真实", "拟像", "仿真", "符号价值"])}，其与{kw1}的关系已经发生了{random_choice(["内爆", "外爆", "逆转", "消解"])}。',

            # 后人类/技术类
            f'随着{random_choice(["技术奇点", "后人类状况", "赛博空间", "数字转向"])}的到来，'
            f'{kw1}与{kw2}的边界正在{random_choice(["模糊", "消解", "重构", "重新编码"])}，'
            f'这对传统的{random_choice(["人本主义", "主体哲学", "意识哲学", "人类中心主义"])}构成了根本性挑战。',

            # 复杂系统类
            f'从{random_choice(["复杂系统理论", "自组织理论", "耗散结构理论", "协同学"])}的观点来看，'
            f'{kw2}的产生是一种{random_choice(["涌现现象", "自组织过程", "相变行为", "临界现象"])}，'
            f'它体现了{kw1}系统在{random_choice(["混沌边缘", "临界点", "分岔点", "远平衡态"])}的特性。',

            # 量子/神秘类
            f'某种意义上，{kw1}与{kw2}的关系就像{random_choice(["量子纠缠", "波粒二象性", "测不准原理", "观察者效应"])}，'
            f'它们{random_choice(["不可分割", "相互建构", "互为因果", "同时存在"])}，'
            f'任何试图将其{random_choice(["分离", "还原", "客观化", "对象化"])}的努力都是徒劳的。',

            # 辩证法类
            f'辩证地看，{kw1}与{kw2}之间存在着{random_choice(["对立统一", "辩证否定", "质量互变", "否定之否定"])}的关系，'
            f'它们在{kw3}的中介下，实现了{random_choice(["螺旋式上升", "矛盾转化", "辩证综合", "具体统一"])}。'
        ]

    # 英文术语混搭类模板生成函数
    def get_english_mixed_sentences():
        return [
            # 理论框架类
            f'基于{random_choice(ENGLISH_TERMS)}理论框架，我们提出了'
            f'{kw1}与{kw2}相互作用的{random_choice(["数学模型", "理论模型", "计算框架", "唯象理论"])}，'
            f'该模型成功解释了{kw3}条件下观察到的各种实验现象。',

            # 方法学类
            f'我们采用了{random_choice(ENGLISH_TERMS)}方法来研究{kw2}的'
            f'{random_choice(["动力学行为", "统计特性", "相变规律", "响应机制"])}，'
            f'结果表明该方法具有{random_choice(["高精度", "高效率", "强鲁棒性", "广泛适用性"])}。',

            # 现象描述类
            f'实验中观察到了典型的{random_choice(ENGLISH_TERMS)}现象，'
            f'其特征与{kw1}的{random_choice(["强度", "频率", "浓度", "温度"])}密切相关，'
            f'这为理解{kw2}的微观机制提供了重要线索。',

            # 机制解释类
            f'我们认为，这一现象的物理机制可以用'
            f'{random_choice(ENGLISH_TERMS)}与{random_choice(ENGLISH_TERMS)}的耦合效应来解释，'
            f'其中{kw3}起着关键的调控作用。',

            # 比较研究类
            f'与传统的{random_choice(["理论", "方法", "模型"])}相比，'
            f'基于{random_choice(ENGLISH_TERMS)}的新方法能够更好地描述{kw2}的'
            f'{random_choice(["非平衡态行为", "非线性特性", "量子效应", "多尺度行为"])}。',

            # 数值模拟类
            f'通过{random_choice(ENGLISH_TERMS)}数值模拟，我们重现了实验中观察到的'
            f'{kw1}与{kw2}的{random_choice(["相关性", "因果关系", "耦合效应", "同步行为"])}，'
            f'模拟结果与实验数据定性一致。',

            # 临界现象类
            f'在临界点附近，系统表现出显著的{random_choice(ENGLISH_TERMS)}特征，'
            f'{kw1}的微小变化会导致{kw2}的{random_choice(["突变", "发散", "标度行为", "普适性"])}，'
            f'这与{random_choice(ENGLISH_TERMS)}理论的预测完全吻合。',

            # 多尺度类
            f'研究发现，{kw2}具有明显的{random_choice(ENGLISH_TERMS)}特性，'
            f'在不同尺度上表现出不同的行为，需要用{random_choice(ENGLISH_TERMS)}方法来统一描述。',

            # 量子/经典对应类
            f'在{random_choice(["经典极限", "热力学极限", "半经典近似", "平均场近似"])}下，'
            f'{random_choice(ENGLISH_TERMS)}效应逐渐消失，系统退化为经典的{kw1}行为，'
            f'这验证了量子-经典对应原理。',

            # 信息论类
            f'从{random_choice(ENGLISH_TERMS)}的角度来看，'
            f'{kw1}与{kw2}之间的{random_choice(["信息传递", "熵交换", "互信息", "信息熵"])}'
            f'可以用来定量描述它们的耦合强度。',

            # 网络/系统类
            f'{kw2}可以被建模为一个复杂的{random_choice(ENGLISH_TERMS)}系统，'
            f'其中{kw1}作为关键节点发挥着重要作用，'
            f'系统的整体行为由{random_choice(["网络拓扑", "节点动力学", "耦合方式", "外部驱动"])}共同决定。',

            # 总结展望类
            f'本研究不仅丰富了{random_choice(ENGLISH_TERMS)}领域的研究内容，'
            f'也为理解{kw1}与{kw2}的关系提供了新的视角，'
            f'未来的研究可以进一步探索{random_choice(ENGLISH_TERMS)}在其中的作用。'
        ]

    # 普通学术废话模板生成函数
    def get_normal_sentences():
        return [
            '这一发现具有重要的理论意义和潜在的应用价值。',
            '进一步的研究正在进行中，相关结果将在后续工作中报道。',
            f'这为我们理解{kw2}的本质提供了新的视角。',
            '相关的深入分析将在下一节详细讨论。',
            '这一结果对相关领域的研究具有一定的参考价值。',
            '更多的实验验证和理论分析仍在进行中。',
            '这些发现为后续研究奠定了基础。',
            '该方法具有一定的通用性和推广价值。',
            '详细的推导过程可参见附录。',
            '这与我们的预期基本一致。',
            '这一结果与已有的文献报道基本一致。',
            '进一步的优化工作正在进行中。',
            '这为相关领域的研究提供了新的思路。',
            '相关的应用研究也在逐步展开。',
            '这一现象值得进一步深入研究。',
            '我们将在后续工作中详细讨论这一问题。',
            '这一结论得到了实验数据的有力支持。',
            '该方法的适用范围还有待进一步扩展。',
            '这些结果具有重要的学术价值和应用前景。',
            '总的来说，本研究取得了预期的成果。'
        ]

    # 结尾句模板生成函数
    def get_endings():
        return [
            # 总结类
            f'综上所述，{section_title}为后续研究奠定了'
            f'{random_choice(["理论基础", "实验基础", "方法基础", "数据基础", "重要基础"])}。',
            f'总的来说，{section_title}的研究对于理解{kw2}的'
            f'{random_choice(["本质", "机制", "行为", "演化"])}具有重要意义。',

            # 过渡类
            '这些结果将在下一章中进行更深入的讨论。',
            '接下来，我们将通过实验来验证这一理论预测。',
            '相关内容将在后续章节中进一步展开。',
            '下一节将详细讨论这一现象的物理机制。',

            # 线索类
            f'这为我们理解{kw1}与{kw2}的关系提供了重要线索。',
            f'这些发现为揭示{kw3}的作用机制打开了新的窗口。',

            # 展望类
            '未来的研究将进一步探索这一现象的深层机制。',
            '更多的实验和理论工作仍有待开展。',

            # 结论类
            '本节的结果为全文的分析提供了重要支撑。',
            f'通过本节的讨论，我们对{section_title}有了更全面的认识。'
        ]

    # 句子类型权重
    type_weights = {
        'theory': 0.18,
        'experiment': 0.15,
        'result': 0.12,
        'mechanism': 0.12,
        'citation': 0.10,
        'crazy': 0.08,
        'chaos': 0.07,
        'english': 0.06,
        'normal': 0.12
    }

    # 根据强度调整胡说八道比例
    if intensity > 0.5:
        type_weights['crazy'] = 0.06 + intensity * 0.1
        type_weights['chaos'] = 0.05 + intensity * 0.05
        type_weights['normal'] = max(0.05, 0.15 - intensity * 0.1)

    # 生成开头句
    openings = get_openings()
    sentences.append(random_choice(openings))

    # 生成中间内容
    n_sentences = random_int(4, 8) + int(intensity * 4)
    type_generators = {
        'theory': get_theory_sentences,
        'experiment': get_experiment_sentences,
        'result': get_result_sentences,
        'mechanism': get_mechanism_sentences,
        'citation': get_citation_sentences,
        'crazy': get_crazy_sentences,
        'chaos': get_chaos_sentences,
        'english': get_english_mixed_sentences,
        'normal': get_normal_sentences
    }

    for _ in range(n_sentences):
        # 加权随机选择句子类型
        r = random.random()
        cumulative = 0.0
        selected_type = 'normal'

        for typ, weight in type_weights.items():
            cumulative += weight
            if r < cumulative:
                selected_type = typ
                break

        # 从对应类型的模板中随机选择
        templates = type_generators[selected_type]()
        template_idx = random_int(0, len(templates) - 1)

        # 尝试避免同一段落中重复使用相同模板（最多尝试3次）
        attempts = 0
        while f'{selected_type}-{template_idx}' in used_templates and attempts < 3:
            template_idx = random_int(0, len(templates) - 1)
            attempts += 1
        used_templates.add(f'{selected_type}-{template_idx}')

        sentences.append(templates[template_idx])

    # 生成结尾句
    endings = get_endings()
    sentences.append(random_choice(endings))

    return ''.join(sentences)


def generate_table(doc, kw1, kw2, kw3, table_num):
    """生成表格

    Args:
        doc: docx的Document对象
        kw1: 关键词1
        kw2: 关键词2
        kw3: 关键词3
        table_num: 表格编号
    """
    n_rows = random_int(4, 7)
    n_cols = random_int(4, 6)

    # 表格类型
    table_types = [
        'response',     # 响应特性表
        'comparison',   # 对比表
        'parameter',    # 参数表
        'statistic',    # 统计表
        'result'        # 结果表
    ]
    table_type = random_choice(table_types)

    # 表格标题模板
    table_title_templates = {
        'response': [
            f'不同条件下{kw2}的{kw1}响应特性',
            f'{kw1}强度对{kw2}性能的影响',
            f'不同{kw3}水平下{kw2}的响应曲线',
            f'{kw1}与{kw2}的剂量-效应关系',
            f'不同实验条件下{kw2}的测量结果'
        ],
        'comparison': [
            f'不同方法在{kw2}研究中的对比分析',
            f'{kw1}与{kw3}对{kw2}影响的比较',
            f'各组实验结果的对比分析',
            f'不同条件下{kw2}性能的综合比较',
            f'本文方法与传统方法的性能对比'
        ],
        'parameter': [
            f'实验样品的主要参数',
            f'{kw2}的特征参数汇总',
            f'不同样品的物理参数对比',
            f'实验所用主要参数列表',
            f'{kw1}的典型参数值'
        ],
        'statistic': [
            f'实验数据的统计分析结果',
            f'各组数据的描述性统计',
            f'{kw2}测量结果的统计特征',
            f'方差分析结果汇总表',
            f'相关性分析结果'
        ],
        'result': [
            f'主要实验结果汇总',
            f'{kw1}作用下{kw2}的变化规律',
            f'不同条件下的实验结果',
            f'{kw3}对{kw2}影响的实验结果',
            f'正交实验结果分析'
        ]
    }

    table_title = random_choice(table_title_templates[table_type])

    # 列标题模板
    col_header_templates = {
        'response': [
            ['样品编号', f'{kw1}强度', '测量值', '误差', f'{kw2}响应'],
            ['实验组', '参数A', '参数B', '平均值', '标准差'],
            ['条件编号', '温度', '压力', f'{kw2}值', '备注'],
            ['测量次数', f'{kw1}值', f'{kw2}值', '相对偏差', '说明']
        ],
        'comparison': [
            ['方法名称', '精度', '效率', '适用范围', '优缺点'],
            ['实验组别', '平均值', '标准差', '显著性', '结论'],
            ['对比项目', '方法A', '方法B', '本文方法', '最优'],
            ['性能指标', '传统方法', '改进方法', '提升幅度', '说明']
        ],
        'parameter': [
            ['参数名称', '符号', '数值', '单位', '备注'],
            ['样品编号', '尺寸', '重量', '纯度', '来源'],
            ['设备名称', '型号', '精度', '量程', '生产厂家'],
            ['物理量', '符号', '数值', '误差', '测量方法']
        ],
        'statistic': [
            ['统计量', '数值', '标准误', '置信区间', '显著性'],
            ['组别', '样本量', '平均值', '标准差', '标准误'],
            ['变量', '均值', '方差', '偏度', '峰度'],
            ['分析项', 'F值', 'p值', '显著性', '结论']
        ],
        'result': [
            ['实验编号', '条件1', '条件2', '结果', '结论'],
            ['序号', f'{kw1}', f'{kw2}', f'{kw3}', '备注'],
            ['实验组', '初始值', '最终值', '变化率', '说明'],
            ['测试项目', '测试结果', '标准值', '偏差', '结论']
        ]
    }

    col_headers = random_choice(col_header_templates[table_type])

    # 行标签模板
    row_label_templates = {
        'response': ['样品A', '样品B', '样品C', '样品D', '样品E', '样品F', '样品G'],
        'comparison': ['方法一', '方法二', '方法三', '方法四', '方法五'],
        'parameter': ['参数1', '参数2', '参数3', '参数4', '参数5', '参数6'],
        'statistic': ['组1', '组2', '组3', '组4', '组5'],
        'result': ['实验1', '实验2', '实验3', '实验4', '实验5', '实验6']
    }

    row_labels = row_label_templates[table_type]

    # 表题
    add_table_caption(doc, f'表{table_num}  {table_title}')

    # 创建表格
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for j in range(n_cols):
        hdr_cells[j].text = col_headers[j % len(col_headers)]

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, size=Pt(CAPTION_SIZE), bold=True)

    # 数据行
    for i in range(n_rows):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = row_labels[i % len(row_labels)]

        for j in range(1, n_cols):
            rand = random.random()
            if rand < 0.3:
                # 数值型
                value = f'{random_float(0.001, 1000, 3)}'
            elif rand < 0.5:
                # 百分比型
                value = f'{random_float(0.1, 99.9, 2)}%'
            elif rand < 0.7:
                # 误差型
                value = f'±{random_float(0.01, 10, 2)}'
            elif rand < 0.85:
                # 整数型
                value = f'{random_int(1, 1000)}'
            else:
                # 文字型
                value = random_choice(['显著', '不显著', '正常', '异常', '合格', '不合格', '优', '良', '中', '差'])

            row_cells[j].text = value

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_font(run, size=Pt(CAPTION_SIZE))

    doc.add_paragraph()  # 空行


# ============== 图表生成函数（Pillow） ==============

def _get_font(size=12):
    """获取字体，优先使用中文字体"""
    font_paths = [
        'C:/Windows/Fonts/simhei.ttf',  # 黑体
        'C:/Windows/Fonts/simsun.ttc',  # 宋体
        'C:/Windows/Fonts/msyh.ttc',    # 微软雅黑
    ]
    for path in font_paths:
        try:
            return ImageFont.truetype(path, size)
        except (IOError, OSError):
            continue
    return ImageFont.load_default()


def _draw_axes(draw, x, y, width, height, x_labels, y_max, y_min=0, font=None):
    """绘制坐标轴和网格"""
    if font is None:
        font = _get_font(10)
    
    # 坐标轴颜色
    axis_color = '#333333'
    grid_color = '#e0e0e0'
    
    # 绘制 Y 轴网格线和刻度
    n_grid = 5
    for i in range(n_grid + 1):
        gy = y + height - int(i * height / n_grid)
        value = y_min + (y_max - y_min) * i / n_grid
        draw.line([(x, gy), (x + width, gy)], fill=grid_color, width=1)
        
        # Y 轴刻度文字
        val_str = f'{value:.1f}' if abs(value) < 10 else f'{value:.0f}'
        draw.text((x - 5, gy), val_str, fill=axis_color, font=font, anchor='rm')
    
    # 绘制 X 轴和 Y 轴
    draw.line([(x, y), (x, y + height)], fill=axis_color, width=2)
    draw.line([(x, y + height), (x + width, y + height)], fill=axis_color, width=2)
    
    # 绘制 X 轴标签
    if x_labels:
        n_labels = len(x_labels)
        for i, label in enumerate(x_labels):
            lx = x + int((i + 0.5) * width / n_labels)
            draw.text((lx, y + height + 5), str(label), fill=axis_color, 
                      font=font, anchor='mt')


def generate_bar_chart(title, x_labels, data_sets, width=600, height=350):
    """生成柱状图"""
    colors = ['#1976d2', '#ff9800', '#4caf50', '#e91e63', '#9c27b0']
    
    # 计算 Y 轴范围
    all_values = []
    for ds in data_sets:
        all_values.extend(ds['data'])
    y_max = max(all_values) * 1.1 if all_values else 10
    y_min = 0
    
    # 创建图像
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    font_title = _get_font(14)
    font_axis = _get_font(10)
    font_legend = _get_font(10)
    
    # 边距
    margin_left = 60
    margin_right = 20
    margin_top = 40
    margin_bottom = 50
    chart_width = width - margin_left - margin_right
    chart_height = height - margin_top - margin_bottom
    
    # 绘制标题
    draw.text((width // 2, 15), title, fill='#333333', font=font_title, anchor='mt')
    
    # 绘制坐标轴
    _draw_axes(draw, margin_left, margin_top, chart_width, chart_height, 
               x_labels, y_max, y_min, font_axis)
    
    # 绘制柱状图
    n_groups = len(x_labels)
    n_series = len(data_sets)
    group_width = chart_width / n_groups
    bar_width = group_width * 0.7 / n_series
    
    for i, ds in enumerate(data_sets):
        color = colors[i % len(colors)]
        for j, value in enumerate(ds['data']):
            bar_x = margin_left + j * group_width + i * bar_width + bar_width * 0.15
            bar_height = int((value - y_min) / (y_max - y_min) * chart_height)
            bar_y = margin_top + chart_height - bar_height
            draw.rectangle([bar_x, bar_y, bar_x + bar_width * 0.85, 
                           margin_top + chart_height], fill=color)
    
    # 绘制图例
    legend_x = margin_left
    legend_y = height - 25
    for i, ds in enumerate(data_sets):
        color = colors[i % len(colors)]
        lx = legend_x + i * 120
        draw.rectangle([lx, legend_y, lx + 12, legend_y + 12], fill=color)
        draw.text((lx + 18, legend_y + 1), ds.get('label', f'系列{i+1}'), 
                  fill='#333333', font=font_legend)
    
    return img


def generate_line_chart(title, x_labels, data_sets, width=600, height=350):
    """生成折线图"""
    colors = ['#1976d2', '#ff9800', '#4caf50', '#e91e63', '#9c27b0']
    
    # 计算 Y 轴范围
    all_values = []
    for ds in data_sets:
        all_values.extend(ds['data'])
    y_max = max(all_values) * 1.1 if all_values else 10
    y_min = min(all_values) * 0.9 if all_values and min(all_values) < 0 else 0
    
    # 创建图像
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    font_title = _get_font(14)
    font_axis = _get_font(10)
    font_legend = _get_font(10)
    
    # 边距
    margin_left = 60
    margin_right = 20
    margin_top = 40
    margin_bottom = 50
    chart_width = width - margin_left - margin_right
    chart_height = height - margin_top - margin_bottom
    
    # 绘制标题
    draw.text((width // 2, 15), title, fill='#333333', font=font_title, anchor='mt')
    
    # 绘制坐标轴
    _draw_axes(draw, margin_left, margin_top, chart_width, chart_height, 
               x_labels, y_max, y_min, font_axis)
    
    # 绘制折线
    n_points = len(x_labels)
    
    for i, ds in enumerate(data_sets):
        color = colors[i % len(colors)]
        points = []
        for j, value in enumerate(ds['data']):
            px = margin_left + int((j + 0.5) * chart_width / n_points)
            py = margin_top + chart_height - int((value - y_min) / (y_max - y_min) * chart_height)
            points.append((px, py))
        
        # 绘制线条
        if len(points) > 1:
            draw.line(points, fill=color, width=2)
        
        # 绘制数据点
        for px, py in points:
            draw.ellipse([px - 3, py - 3, px + 3, py + 3], fill=color, outline='white', width=1)
    
    # 绘制图例
    legend_x = margin_left
    legend_y = height - 25
    for i, ds in enumerate(data_sets):
        color = colors[i % len(colors)]
        lx = legend_x + i * 120
        draw.line([lx, legend_y + 6, lx + 12, legend_y + 6], fill=color, width=2)
        draw.ellipse([lx + 3, legend_y + 3, lx + 9, legend_y + 9], fill=color)
        draw.text((lx + 18, legend_y + 1), ds.get('label', f'系列{i+1}'), 
                  fill='#333333', font=font_legend)
    
    return img


def generate_pie_chart(title, data, width=500, height=350):
    """生成饼图"""
    colors = ['#1976d2', '#ff9800', '#4caf50', '#e91e63', '#9c27b0', 
              '#00bcd4', '#ffeb3b', '#795548']
    
    # 创建图像
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    font_title = _get_font(14)
    font_label = _get_font(10)
    
    # 边距
    margin_left = 20
    margin_right = 180
    margin_top = 40
    margin_bottom = 20
    chart_size = min(width - margin_left - margin_right, height - margin_top - margin_bottom)
    chart_x = margin_left
    chart_y = margin_top
    
    # 绘制标题
    draw.text((width // 2, 15), title, fill='#333333', font=font_title, anchor='mt')
    
    # 计算总值
    total = sum(item['value'] for item in data)
    
    # 绘制饼图
    start_angle = -90  # 从 12 点钟方向开始
    bbox = [chart_x, chart_y, chart_x + chart_size, chart_y + chart_size]
    
    for i, item in enumerate(data):
        angle = item['value'] / total * 360
        end_angle = start_angle + angle
        
        # 绘制扇形
        color = colors[i % len(colors)]
        draw.pieslice(bbox, start_angle, end_angle, fill=color, outline='white', width=2)
        
        start_angle = end_angle
    
    # 绘制图例
    legend_x = width - margin_right + 10
    legend_y = margin_top
    for i, item in enumerate(data):
        color = colors[i % len(colors)]
        ly = legend_y + i * 25
        draw.rectangle([legend_x, ly, legend_x + 12, ly + 12], fill=color)
        
        percentage = item['value'] / total * 100
        label = f"{item['label']} ({percentage:.1f}%)"
        draw.text((legend_x + 18, ly - 1), label, fill='#333333', font=font_label)
    
    return img


def generate_area_chart(title, x_labels, data_sets, width=600, height=350):
    """生成面积图"""
    colors = ['#1976d2', '#ff9800', '#4caf50', '#e91e63', '#9c27b0']
    
    # 计算 Y 轴范围
    all_values = []
    for ds in data_sets:
        all_values.extend(ds['data'])
    y_max = max(all_values) * 1.1 if all_values else 10
    y_min = 0
    
    # 创建图像
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    font_title = _get_font(14)
    font_axis = _get_font(10)
    font_legend = _get_font(10)
    
    # 边距
    margin_left = 60
    margin_right = 20
    margin_top = 40
    margin_bottom = 50
    chart_width = width - margin_left - margin_right
    chart_height = height - margin_top - margin_bottom
    
    # 绘制标题
    draw.text((width // 2, 15), title, fill='#333333', font=font_title, anchor='mt')
    
    # 绘制坐标轴
    _draw_axes(draw, margin_left, margin_top, chart_width, chart_height, 
               x_labels, y_max, y_min, font_axis)
    
    # 绘制面积图（从后往前绘制，避免遮挡）
    n_points = len(x_labels)
    
    for i in range(len(data_sets) - 1, -1, -1):
        ds = data_sets[i]
        color = colors[i % len(colors)]
        
        # 计算半透明颜色
        # Pillow 不直接支持半透明填充，我们用浅色代替
        r = int(color[1:3], 16)
        g = int(color[3:5], 16)
        b = int(color[5:7], 16)
        light_color = f'#{int(r + (255-r)*0.6):02x}{int(g + (255-g)*0.6):02x}{int(b + (255-b)*0.6):02x}'
        
        points = []
        for j, value in enumerate(ds['data']):
            px = margin_left + int((j + 0.5) * chart_width / n_points)
            py = margin_top + chart_height - int((value - y_min) / (y_max - y_min) * chart_height)
            points.append((px, py))
        
        # 添加底部点形成闭合区域
        points.append((margin_left + chart_width, margin_top + chart_height))
        points.append((margin_left, margin_top + chart_height))
        
        # 绘制填充区域
        draw.polygon(points, fill=light_color)
        
        # 绘制顶部线条
        line_points = points[:-2]
        if len(line_points) > 1:
            draw.line(line_points, fill=color, width=2)
    
    # 绘制图例
    legend_x = margin_left
    legend_y = height - 25
    for i, ds in enumerate(data_sets):
        color = colors[i % len(colors)]
        lx = legend_x + i * 120
        r = int(color[1:3], 16)
        g = int(color[3:5], 16)
        b = int(color[5:7], 16)
        light_color = f'#{int(r + (255-r)*0.6):02x}{int(g + (255-g)*0.6):02x}{int(b + (255-b)*0.6):02x}'
        draw.rectangle([lx, legend_y, lx + 12, legend_y + 12], fill=light_color, outline=color)
        draw.text((lx + 18, legend_y + 1), ds.get('label', f'系列{i+1}'), 
                  fill='#333333', font=font_legend)
    
    return img


def generate_scatter_chart(title, x_label, y_label, data_sets, width=600, height=350):
    """生成散点图"""
    colors = ['#1976d2', '#ff9800', '#4caf50', '#e91e63', '#9c27b0']
    
    # 计算轴范围
    all_x = []
    all_y = []
    for ds in data_sets:
        for point in ds['data']:
            all_x.append(point[0])
            all_y.append(point[1])
    
    x_max = max(all_x) * 1.1 if all_x else 10
    x_min = min(all_x) * 0.9 if all_x and min(all_x) < 0 else 0
    y_max = max(all_y) * 1.1 if all_y else 10
    y_min = min(all_y) * 0.9 if all_y and min(all_y) < 0 else 0
    
    # 创建图像
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    font_title = _get_font(14)
    font_axis = _get_font(10)
    font_legend = _get_font(10)
    
    # 边距
    margin_left = 60
    margin_right = 20
    margin_top = 40
    margin_bottom = 60
    chart_width = width - margin_left - margin_right
    chart_height = height - margin_top - margin_bottom
    
    # 绘制标题
    draw.text((width // 2, 15), title, fill='#333333', font=font_title, anchor='mt')
    
    # 绘制坐标轴和网格
    _draw_axes(draw, margin_left, margin_top, chart_width, chart_height, 
               None, y_max, y_min, font_axis)
    
    # X 轴刻度
    n_grid = 5
    for i in range(n_grid + 1):
        gx = margin_left + int(i * chart_width / n_grid)
        value = x_min + (x_max - x_min) * i / n_grid
        val_str = f'{value:.1f}' if abs(value) < 10 else f'{value:.0f}'
        draw.text((gx, margin_top + chart_height + 5), val_str, 
                  fill='#333333', font=font_axis, anchor='mt')
    
    # 轴标签
    draw.text((width // 2, height - 15), x_label, fill='#333333', 
              font=font_axis, anchor='mt')
    draw.text((15, height // 2), y_label, fill='#333333', 
              font=font_axis, anchor='mm')
    
    # 绘制散点
    for i, ds in enumerate(data_sets):
        color = colors[i % len(colors)]
        for point in ds['data']:
            px = margin_left + int((point[0] - x_min) / (x_max - x_min) * chart_width)
            py = margin_top + chart_height - int((point[1] - y_min) / (y_max - y_min) * chart_height)
            draw.ellipse([px - 4, py - 4, px + 4, py + 4], fill=color, outline='white', width=1)
    
    # 绘制图例
    legend_x = margin_left
    legend_y = height - 40
    for i, ds in enumerate(data_sets):
        color = colors[i % len(colors)]
        lx = legend_x + i * 120
        draw.ellipse([lx, legend_y + 2, lx + 10, legend_y + 12], fill=color)
        draw.text((lx + 16, legend_y + 1), ds.get('label', f'系列{i+1}'), 
                  fill='#333333', font=font_legend)
    
    return img


def generate_figure(doc, kw1, kw2, fig_num):
    """生成图

    Args:
        doc: docx的Document对象
        kw1: 关键词1
        kw2: 关键词2
        fig_num: 图编号
    """
    # 随机选择图表类型
    chart_types = ['bar', 'line', 'pie', 'area', 'scatter']
    chart_type = random.choice(chart_types)
    
    # 生成随机数据
    n_categories = random_int(4, 7)
    n_series = random_int(2, 4)
    
    # X 轴标签
    x_labels = [f'{random_choice(["实验组", "对照组", "样本", "组别"])}{i+1}' 
                for i in range(n_categories)]
    
    # 数据集
    data_sets = []
    series_names = [f'{kw1}组', f'{kw2}组', '对照组', '实验组']
    for i in range(n_series):
        data = [random_float(10, 100, 1) for _ in range(n_categories)]
        data_sets.append({
            'label': series_names[i % len(series_names)],
            'data': data
        })
    
    # 图标题
    figure_titles = {
        'bar': f'不同条件下{kw1}与{kw2}的对比分析',
        'line': f'{kw1}与{kw2}的变化趋势图',
        'pie': f'{kw1}的占比分布',
        'area': f'{kw1}与{kw2}的面积对比图',
        'scatter': f'{kw1}与{kw2}的相关性散点图'
    }
    title = figure_titles[chart_type]
    
    # 生成图表
    if chart_type == 'bar':
        img = generate_bar_chart(title, x_labels, data_sets)
    elif chart_type == 'line':
        img = generate_line_chart(title, x_labels, data_sets)
    elif chart_type == 'pie':
        # 饼图数据
        pie_data = []
        for i in range(n_categories):
            pie_data.append({
                'label': x_labels[i],
                'value': random_float(10, 100, 1)
            })
        img = generate_pie_chart(title, pie_data)
    elif chart_type == 'area':
        img = generate_area_chart(title, x_labels, data_sets)
    elif chart_type == 'scatter':
        # 散点图数据
        scatter_data = []
        for i in range(n_series):
            points = []
            n_points = random_int(15, 30)
            for _ in range(n_points):
                x = random_float(0, 100, 1)
                y = random_float(0, 100, 1)
                points.append((x, y))
            scatter_data.append({
                'label': series_names[i % len(series_names)],
                'data': points
            })
        img = generate_scatter_chart(title, kw1, kw2, scatter_data)
    
    # 保存图片到临时文件
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
        tmp_path = tmp.name
        img.save(tmp_path, 'PNG')
    
    # 插入图片到 Word 文档
    fig_p = doc.add_paragraph()
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_pf = fig_p.paragraph_format
    fig_pf.space_before = Pt(12)
    fig_pf.space_after = Pt(6)
    
    run = fig_p.add_run()
    run.add_picture(tmp_path, width=Inches(5.5))
    
    # 删除临时文件
    os.unlink(tmp_path)
    
    # 图题
    add_figure_caption(doc, f'图{fig_num}  {title}')


def generate_formula(doc, formula_num):
    """生成公式

    Args:
        doc: docx的Document对象
        formula_num: 公式编号
    """
    formula = random_choice(FORMULA_TEMPLATES)
    add_formula(doc, f'{formula}    ({formula_num})')


def generate_chinese_reference(ref_num):
    """生成单条中文参考文献

    Args:
        ref_num: 参考文献编号

    Returns:
        参考文献文本
    """
    is_journal = random.random() < 0.7

    if is_journal:
        n_authors = random_int(1, 5)
        authors = random.sample(AUTHOR_NAMES, min(n_authors, len(AUTHOR_NAMES)))
        author_str = '，'.join(authors)

        # 生成中文标题
        title_templates = [
            f'基于{random_choice(ENGLISH_TERMS)}的{random_choice(CHAOS_WORDS)}研究',
            f'{random_choice(CHAOS_WORDS)}视角下的{random_choice(PHYSICAL_QUANTITIES)}分析',
            f'论{random_choice(CHAOS_WORDS)}与{random_choice(PHYSICAL_QUANTITIES)}的辩证关系',
            f'{random_choice(ENGLISH_TERMS)}理论及其在{random_choice(CHAOS_WORDS)}中的应用',
            f'{random_choice(CHAOS_WORDS)}的{random_choice(ENGLISH_TERMS)}机制研究',
            f'关于{random_choice(PHYSICAL_QUANTITIES)}的{random_choice(CHAOS_WORDS)}思考',
            f'{random_choice(ENGLISH_TERMS)}视域中的{random_choice(CHAOS_WORDS)}问题',
            f'{random_choice(CHAOS_WORDS)}与{random_choice(CHAOS_WORDS)}的{random_choice(ENGLISH_TERMS)}分析',
            f'从{random_choice(CHAOS_WORDS)}到{random_choice(CHAOS_WORDS)}：{random_choice(ENGLISH_TERMS)}的路径'
        ]
        title = random_choice(title_templates)

        journal = random_choice(JOURNALS)
        year = random_int(1985, 2025)
        volume = random_int(1, 150)
        issue = random_int(1, 12)
        start_page = random_int(1, 800)
        end_page = start_page + random_int(5, 50)

        return (f'[{ref_num}] {author_str}. {title}[J]. {journal}, '
                f'{year}, {volume}({issue}): {start_page}-{end_page}.')
    else:
        n_authors = random_int(1, 3)
        authors = random.sample(AUTHOR_NAMES, min(n_authors, len(AUTHOR_NAMES)))
        author_str = '，'.join(authors)

        title_templates = [
            f'{random_choice(CHAOS_WORDS)}导论',
            f'{random_choice(ENGLISH_TERMS)}原理与应用',
            f'{random_choice(CHAOS_WORDS)}研究方法',
            f'现代{random_choice(PHYSICAL_QUANTITIES)}学',
            f'{random_choice(CHAOS_WORDS)}与{random_choice(ENGLISH_TERMS)}'
        ]
        title = random_choice(title_templates)

        publisher_city = random_choice([
            '北京', '上海', '南京', '武汉', '广州',
            '成都', '西安', '杭州', '深圳', '哈尔滨'
        ])
        publishers = [
            '胡说八道出版社', '混沌理论出版社', '量子早餐出版社',
            '伪科学出版社', '虚构学术出版社', '抽象研究出版社',
            '非线性出版社', '统计力学出版社', '膜宇宙出版社',
            '弦论出版社'
        ]
        publisher = random_choice(publishers)

        year = random_int(1970, 2024)

        return (f'[{ref_num}] {author_str}. {title}[M]. '
                f'{publisher_city}: {publisher}, {year}.')


def generate_english_reference(ref_num):
    """生成单条英文参考文献

    Args:
        ref_num: 参考文献编号

    Returns:
        参考文献文本
    """
    is_journal = random.random() < 0.7

    if is_journal:
        n_authors = random_int(1, 5)
        authors = []
        for _ in range(n_authors):
            surname = ''.join(random.choices(string.ascii_uppercase, k=1)) + \
                      ''.join(random.choices(string.ascii_lowercase,
                                             k=random.randint(3, 10)))
            initials = ''.join(random.choices(string.ascii_uppercase,
                                              k=random.randint(1, 2))) + '.'
            authors.append(f'{surname} {initials}')

        if n_authors > 2:
            author_str = ', '.join(authors[:-1]) + ', & ' + authors[-1]
        elif n_authors == 2:
            author_str = authors[0] + ' & ' + authors[1]
        else:
            author_str = authors[0]

        # 生成英文标题（使用 ENGLISH_TERMS）
        title_templates = [
            f'A study of {random_choice(ENGLISH_TERMS)} in {random_choice(ENGLISH_TERMS)}',
            f'{random_choice(ENGLISH_TERMS)} and its applications to {random_choice(ENGLISH_TERMS)}',
            f'On the {random_choice(ENGLISH_TERMS)} of {random_choice(ENGLISH_TERMS)}',
            f'{random_choice(ENGLISH_TERMS)} theory and {random_choice(ENGLISH_TERMS)} phenomena',
            f'The {random_choice(ENGLISH_TERMS)} mechanism of {random_choice(ENGLISH_TERMS)}',
            f'{random_choice(ENGLISH_TERMS)}: A {random_choice(ENGLISH_TERMS)} perspective',
            f'{random_choice(ENGLISH_TERMS)} analysis of {random_choice(ENGLISH_TERMS)}',
            f'{random_choice(ENGLISH_TERMS)} and {random_choice(ENGLISH_TERMS)}: A review'
        ]
        title = random_choice(title_templates)

        journal = random_choice(ENGLISH_JOURNALS)
        year = random_int(1985, 2025)
        volume = random_int(1, 150)
        issue = random_int(1, 12)
        start_page = random_int(1, 800)
        end_page = start_page + random_int(5, 50)

        return (f'[{ref_num}] {author_str} ({year}). {title}. {journal}, '
                f'{volume}({issue}), {start_page}-{end_page}.')
    else:
        n_authors = random_int(1, 3)
        authors = []
        for _ in range(n_authors):
            surname = ''.join(random.choices(string.ascii_uppercase, k=1)) + \
                      ''.join(random.choices(string.ascii_lowercase,
                                             k=random.randint(3, 10)))
            initials = ''.join(random.choices(string.ascii_uppercase,
                                              k=random.randint(1, 2))) + '.'
            authors.append(f'{surname} {initials}')

        author_str = ', '.join(authors)

        title_templates = [
            f'Introduction to {random_choice(ENGLISH_TERMS)}',
            f'Principles of {random_choice(ENGLISH_TERMS)}',
            f'{random_choice(ENGLISH_TERMS)} and {random_choice(ENGLISH_TERMS)}',
            f'Advanced {random_choice(ENGLISH_TERMS)}',
            f'Foundations of {random_choice(ENGLISH_TERMS)}'
        ]
        title = random_choice(title_templates)

        publisher_city = random_choice([
            'New York', 'London', 'Berlin', 'Tokyo', 'Paris',
            'Cambridge', 'Oxford', 'Chicago', 'Amsterdam', 'Singapore'
        ])
        publisher_words = random_int(2, 4)
        publisher = ' '.join([
            ''.join(random.choices(string.ascii_uppercase + string.ascii_lowercase,
                                   k=random.randint(3, 12)))
            for _ in range(publisher_words)
        ])
        publisher = publisher.title() + ' ' + random_choice([
            'Press', 'Publishers', 'Books', 'Publications'
        ])

        year = random_int(1970, 2024)

        return (f'[{ref_num}] {author_str} ({year}). {title}. '
                f'{publisher_city}: {publisher}.')


def generate_thanks():
    """生成致谢内容

    Returns:
        致谢段落列表
    """
    n_thanks = random_int(4, 8)
    thanks = []

    # 虚构的致谢对象（动物/植物类）
    thanks_animals_plants = [
        '实验室的量子猫',
        '桌上的混沌多肉',
        '鱼缸里的熵增金鱼',
        '窗台上的薛定谔绿萝',
        '键盘上的随机按键精灵',
        '鼠标垫下的隐藏维度生物',
        '显示器里的像素小精灵'
    ]

    # 虚构的致谢对象（服务类）
    thanks_services = [
        '楼下的量子早餐摊',
        '第四维度的外卖小哥',
        '平行宇宙的咖啡店',
        '暗物质便利店',
        '虫洞快递站',
        '熵增保洁公司',
        '时间循环物业公司'
    ]

    # 虚构的致谢对象（机构类）
    thanks_institutions = [
        '国际胡说八道研究基金会',
        '虚拟科学研究基金',
        '不知名的跨维度审稿人',
        '某神秘混沌组织',
        '平行宇宙研究院',
        '量子早餐动力学实验室',
        '熵增爱好者协会',
        '暗物质研究中心',
        '虫洞交通管理局',
        '时间旅行爱好者俱乐部'
    ]

    # 虚构的致谢对象（理论类）
    thanks_theories = [
        '量子力学的不确定性',
        '热力学第二定律的仁慈',
        '混沌理论的蝴蝶',
        '墨菲定律的眷顾',
        '测不准原理的宽容',
        '熵增定律的缓慢',
        '蝴蝶效应的翅膀',
        '平行宇宙的另一种可能'
    ]

    # 虚构的致谢对象（牺牲品类）
    thanks_sacrifices = [
        '掉的第N根头发',
        '喝掉的第N杯量子咖啡',
        '吃掉的第N份叠加态外卖',
        '坏掉的第N个键盘',
        '报废的第N个鼠标',
        '用光的第N根笔芯',
        '熬的第N个夜',
        '第N次想放弃的念头'
    ]

    # 虚构的致谢对象（人格类）
    thanks_personalities = [
        '平行宇宙的另一个我',
        '未来的自己发来的警告',
        '过去的自己埋下的伏笔',
        '想象中的第N个审稿人',
        '内心的小恶魔',
        '脑海中的第N个声音',
        '梦中的神秘导师',
        '潜意识里的混沌因子'
    ]

    # 虚构的致谢对象（网络类）
    thanks_internet = [
        '量子互联网',
        '混沌搜索引擎',
        '胡说八道百科',
        '跨维度问答社区',
        '高维视频网站',
        '暗网社交平台',
        '虫洞论坛',
        '时间循环留言板'
    ]

    # 虚构的致谢对象（家具类）
    thanks_furniture = [
        '我的量子床',
        '我的叠加态枕头',
        '我的熵增被子',
        '我的混沌沙发',
        '我的时间循环椅子',
        '我的虫洞毯子',
        '我的梦境床垫'
    ]

    # 虚构的致谢对象（物理类）
    thanks_physics = [
        '暗能量',
        '暗物质',
        '希格斯场',
        '量子涨落',
        '真空极化',
        '中微子振荡',
        '夸克禁闭',
        '引力波',
        '宇宙微波背景辐射',
        '黑洞蒸发'
    ]

    # 虚构的致谢对象（读者类）
    thanks_readers = [
        '正在读这篇论文的你',
        '跨维度的审稿人',
        '平行宇宙的读者',
        '外星评委老师',
        '高维空间的导师',
        '时间旅行者',
        '量子观察者'
    ]

    # 合并所有致谢对象
    all_thanks_people = (
        thanks_animals_plants + thanks_services + thanks_institutions +
        thanks_theories + thanks_sacrifices + thanks_personalities +
        thanks_internet + thanks_furniture + thanks_physics + thanks_readers
    )

    # 虚构的致谢物品
    thanks_things = [
        '量子咖啡',
        '混沌奶茶',
        '熵增零食',
        '叠加态外卖',
        '时间压缩饼干',
        '虫洞快递',
        '暗物质能量棒',
        '薛定谔的便当',
        '纠缠态薯片',
        '退相干可乐'
    ]

    thanks_templates = [
        '感谢{person}，没有他/她我早就写完了。',
        '感谢{person}在百忙之中对本文的胡说八道提供了宝贵的意见。',
        '感谢{person}的{thing}，为本文提供了重要的灵感来源。',
        '感谢{person}，虽然他/她完全不知道这篇论文的存在。',
        '感谢{person}在我写论文时送来的{thing}，支撑我完成了这项研究。',
        '特别感谢{person}，是他/她让我明白了什么是真正的胡说八道。',
        '感谢{person}的{thing}，没有它我根本无法坚持到最后。',
        '感谢{person}，他/她的存在本身就是对这项研究最大的支持。',
        '感谢{person}在精神上给予我的支持，虽然我完全没感受到。',
        '感谢{person}提供的{thing}，为本研究提供了重要的物质基础。',
        '感谢{person}在我最困难的时候对我说"加油"，虽然并没有什么用。',
        '感谢{person}，是他/她教会了我，论文写不出来的时候就多喝奶茶。',
        '感谢{person}，每次我想放弃的时候，一想到他/她我就不想写了。',
        '感谢{person}，要不是他/她催我，我还能再拖半年。',
        '感谢{person}，这篇论文的每一个字都与他/她无关。',
        '感谢{person}的{thing}，它陪伴我度过了无数个写不出论文的夜晚。',
        '感谢{person}，虽然他/她什么都没做，但还是要感谢一下。',
        '感谢{person}，是他/她让我知道了，原来论文还可以这么写。'
    ]

    for _ in range(n_thanks):
        template = random_choice(thanks_templates)
        person = random_choice(all_thanks_people)
        thing = random_choice(thanks_things)
        thanks.append(template.format(person=person, thing=thing))

    # 基金项目
    fund_nums = [
        f'国家自然科学基金（No. {random_int(10000000, 99999999)}）',
        f'胡说八道重点研发计划（No. BS{random_int(2020, 2025)}{random_int(100, 999)}）',
        f'莫名其妙青年基金（No. MQ{random_int(2020, 2025)}{random_int(10, 999)}）',
        f'量子烹饪技术专项（No. QC{random_int(2020, 2025)}{random_int(100, 999)}）',
        f'混沌美食研究项目（No. CM{random_int(2020, 2025)}{random_int(100, 999)}）',
        f'膜宇宙早餐工程（No. MB{random_int(2020, 2025)}{random_int(100, 999)}）',
        f'暗袜子探测计划（No. AW{random_int(2020, 2025)}{random_int(100, 999)}）',
        f'跨维度交流合作项目（No. KD{random_int(2020, 2025)}{random_int(100, 999)}）',
        f'平行宇宙联合研究基金（No. PY{random_int(2020, 2025)}{random_int(100, 999)}）',
        f'时间旅行研究专项（No. SJ{random_int(2020, 2025)}{random_int(100, 999)}）'
    ]
    random.shuffle(fund_nums)

    thanks.append(
        f'本研究得到了{"、".join(fund_nums[:random_int(2, 4)])}的资助。'
    )
    thanks.append('当然，这些基金号都是编的，请勿当真。')

    return thanks


# ============== 主生成函数 ==============

def generate_paper(keywords=None, paper_type='journal', level=5,
                   length='medium', output_path=None):
    """生成胡说八道论文

    Args:
        keywords: 关键词列表（0-3个）
        paper_type: 论文类型（journal/master/doctor/conference）
        level: 胡说八道程度（1-10）
        length: 篇幅（short/medium/long）
        output_path: 输出文件路径

    Returns:
        输出文件路径
    """
    # 处理关键词
    kw1, kw2, kw3 = process_keywords(keywords)
    print(f'关键词：{kw1}、{kw2}、{kw3}')

    # 获取章节结构
    if paper_type not in CHAPTER_STRUCTURES:
        paper_type = 'journal'
    chapters = CHAPTER_STRUCTURES[paper_type]
    n_chapters = len(chapters)

    # 根据篇幅调整段落数量系数
    length_multipliers = {
        'short': 0.7,
        'medium': 1.0,
        'long': 1.5
    }
    length_mult = length_multipliers.get(length, 1.0)

    # 生成标题
    title = generate_title(kw1, kw2, kw3)
    print(f'论文标题：{title}')

    # 生成作者和机构
    authors, departments = generate_authors_and_depts()
    print(f'作者：{"、".join(authors)}')

    # 创建文档
    doc = Document()
    sec0 = doc.sections[0]
    
    # 设置默认字体样式
    style = doc.styles['Normal']
    style.font.name = FONT_BODY_EN
    style.font.size = Pt(BODY_FONT_SIZE)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY_CN)
    style.element.rPr.rFonts.set(qn('w:ascii'), FONT_BODY_EN)
    style.element.rPr.rFonts.set(qn('w:hAnsi'), FONT_BODY_EN)
    style.element.rPr.rFonts.set(qn('w:cs'), FONT_BODY_EN)
    
    # 设置默认段落格式
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0.74)
    
    setup_page(sec0)

    # ========== 标题 ==========
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_pf = title_p.paragraph_format
    title_pf.space_before = Pt(36)
    title_pf.space_after = Pt(24)
    title_pf.line_spacing = 1.5
    set_font(title_p.add_run(title),
             FONT_HEADING_CN, FONT_HEADING_EN,
             Pt(18), bold=True)

    # 作者信息
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_pf = author_p.paragraph_format
    author_pf.space_after = Pt(12)
    author_text = '  '.join([f'{a}{i+1}' for i, a in enumerate(authors)])
    set_font(author_p.add_run(author_text))

    # 单位
    dept_p = doc.add_paragraph()
    dept_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dept_pf = dept_p.paragraph_format
    dept_pf.space_after = Pt(24)
    dept_lines = []
    for i, dept in enumerate(departments):
        dept_lines.append(f'{i+1}. {dept}')
    set_font(dept_p.add_run('\n'.join(dept_lines)), size=Pt(10.5))

    # ========== 摘要 ==========
    add_heading(doc, '摘  要', 1)

    abstract_text = generate_abstract(kw1, kw2, kw3, level)
    add_body(doc, abstract_text)

    # 关键词
    keywords_p = doc.add_paragraph()
    keywords_pf = keywords_p.paragraph_format
    keywords_pf.space_before = Pt(12)
    keywords_pf.space_after = Pt(24)
    keywords_pf.first_line_indent = Cm(0.74)
    set_font(keywords_p.add_run('关键词：'), bold=True)

    keywords_list = generate_keywords(kw1, kw2, kw3)
    set_font(keywords_p.add_run('；'.join(keywords_list)))

    # ========== 正文章节 ==========
    # 图表计数器
    fig_count = 0
    table_count = 0
    formula_count = 0

    for ch_idx, chapter in enumerate(chapters):
        chapter_title = chapter['title']
        sections = chapter['sections']

        # 章节标题
        if paper_type == 'conference':
            heading_text = f'{ch_idx + 1}. {chapter_title}'
        else:
            heading_text = f'第{ch_idx + 1}章  {chapter_title}'
        add_heading(doc, heading_text, 1)

        for sec_idx, section_title in enumerate(sections):
            # 小节标题
            if paper_type == 'conference':
                sec_heading = f'{ch_idx + 1}.{sec_idx + 1}  {section_title}'
            else:
                sec_heading = f'{ch_idx + 1}.{sec_idx + 1}  {section_title}'
            add_heading(doc, sec_heading, 2)

            # 生成段落
            n_paragraphs = max(1, int(random_int(2, 5) * length_mult))
            for _ in range(n_paragraphs):
                para_text = generate_paragraph(
                    kw1, kw2, kw3, section_title, chapter_title, level
                )
                add_body(doc, para_text)

            # 随机插入公式、图表
            rand_val = random.random()

            if rand_val < 0.20:
                # 插入公式
                formula_count += 1
                generate_formula(doc, formula_count)
                add_body(doc,
                         f'式（{formula_count}）中，各符号的物理意义如前文所述。'
                         f'该式描述了{kw1}与{kw2}之间的定量关系。')

            elif rand_val < 0.35:
                # 插入表格
                table_count += 1
                generate_table(doc, kw1, kw2, kw3, table_count)
                add_body(doc,
                         f'从表{table_count}可以看出，'
                         f'{random_choice(["随着参数的变化", "不同组之间", "实验结果表明", "数据显示"])}'
                         f'{kw1}对{kw2}具有显著的调控作用。')

            elif rand_val < 0.50:
                # 插入图
                fig_count += 1
                generate_figure(doc, kw1, kw2, fig_count)
                add_body(doc,
                         f'如图{fig_count}所示，{kw1}与{kw2}之间呈现出'
                         f'{random_choice(["明显的正相关", "复杂的非线性", "周期性变化", "先增后减的"])}关系。')

    # ========== 结论 ==========
    # （结论已经包含在章节结构中，不需要额外添加）

    # ========== 参考文献 ==========
    add_heading(doc, '参考文献', 1)

    n_refs = random_int(15, 30)
    n_chinese = int(n_refs * 0.6)  # 60% 中文
    n_english = n_refs - n_chinese  # 40% 英文

    # 中文参考文献
    for i in range(n_chinese):
        ref = generate_chinese_reference(i + 1)
        add_reference(doc, ref)

    # 英文参考文献
    for i in range(n_english):
        ref = generate_english_reference(n_chinese + i + 1)
        add_reference(doc, ref)

    print(f'参考文献：{n_refs}篇（中文{n_chinese}篇，英文{n_english}篇）')

    # ========== 致谢 ==========
    add_heading(doc, '致  谢', 1)

    thanks_list = generate_thanks()
    for t in thanks_list:
        add_body(doc, t)

    # 保存文档
    if output_path is None:
        # 清理标题中的特殊字符作为文件名
        safe_title = ''.join(
            c for c in title
            if c.isalnum() or c in ['-', '_', ' ', '的', '与', '及', '在', '中', '对', '和', '：']
        )
        safe_title = safe_title[:40]  # 限制长度
        output_path = f'{safe_title}.docx'

    doc.save(output_path)
    print(f'\n论文已生成：{output_path}')
    print(f'图表数量：图{fig_count}张，表{table_count}张，公式{formula_count}个')
    print(f'章节数量：{n_chapters}章')

    return output_path


# ============== 命令行接口 ==============

def main():
    """主函数：命令行接口"""
    parser = argparse.ArgumentParser(
        description='胡说八道论文生成器 - 学术严谨版',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  python paper_generator.py                          # 完全随机生成
  python paper_generator.py 量子纠缠 奶茶 热力学     # 指定3个关键词
  python paper_generator.py 螺蛳粉 --type master     # 指定论文类型
  python paper_generator.py 烤红薯 --level 9         # 胡说八道程度拉满
  python paper_generator.py 火锅 --length long       # 长篇论文
        """
    )

    parser.add_argument(
        'keywords',
        nargs='*',
        help='关键词（0-3个，不填则完全随机）'
    )

    parser.add_argument(
        '--type', '-t',
        choices=['journal', 'master', 'doctor', 'conference'],
        default='journal',
        help='论文类型：journal(学术期刊论文)、master(硕士学位论文)、'
             'doctor(博士学位论文)、conference(会议论文)，默认journal'
    )

    parser.add_argument(
        '--level', '-l',
        type=int,
        default=5,
        choices=range(1, 11),
        metavar='1-10',
        help='胡说八道程度（1-10），默认5'
    )

    parser.add_argument(
        '--length', '-L',
        choices=['short', 'medium', 'long'],
        default='medium',
        help='论文篇幅：short(短篇)、medium(中篇)、long(长篇)，默认medium'
    )

    parser.add_argument(
        '--output', '-o',
        help='输出文件路径'
    )

    parser.add_argument(
        '--seed', '-s',
        type=int,
        help='随机种子（用于复现结果）'
    )

    args = parser.parse_args()

    # 设置随机种子
    if args.seed is not None:
        random.seed(args.seed)

    # 生成论文
    print('=' * 60)
    print('胡说八道论文生成器 - 学术严谨版')
    print('=' * 60)
    print()

    generate_paper(
        keywords=args.keywords if args.keywords else None,
        paper_type=args.type,
        level=args.level,
        length=args.length,
        output_path=args.output
    )

    print()
    print('生成完成！')


if __name__ == '__main__':
    main()
